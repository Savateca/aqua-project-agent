from __future__ import annotations

import subprocess
from pathlib import Path
from typing import List, Tuple

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from .report import build_profile_sections


def qn(tag: str) -> str:
    prefix, tagroot = tag.split(':')
    uri_map = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    return f'{{{uri_map[prefix]}}}{tagroot}'


def brl(value: float | None) -> str:
    if value is None:
        return "-"
    return f"R$ {value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def num(value: float | None, digits: int = 2) -> str:
    if value is None:
        return "-"
    return f"{value:,.{digits}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def configure_page(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)


def add_cover(doc: Document, title: str, author: str, subtitle: str = "Relatório técnico-econômico premium"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(120)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(24)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_before = Pt(20)
    p2.add_run(subtitle).font.size = Pt(13)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.space_before = Pt(180)
    p3.add_run(author).font.size = Pt(12)

    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_toc(doc: Document):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.add_run("Sumário")
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)
    doc.add_page_break()


def add_bullets(doc: Document, items: List[str]):
    for item in items:
        if item:
            doc.add_paragraph(item, style="List Bullet")


def add_chart_image(doc: Document, title: str, image_path: Path, caption: str | None = None):
    doc.add_paragraph(style="Heading 2").add_run(title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(6.5))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)


def _save_line_chart(df: pd.DataFrame, x: str, y: str, title: str, ylabel: str, path: Path):
    fig, ax = plt.subplots(figsize=(8.5, 4.6))
    ax.plot(df[x], df[y], marker="o", linewidth=2.5)
    ax.set_title(title)
    ax.set_xlabel("Dia")
    ax.set_ylabel(ylabel)
    ax.grid(True, alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def _save_bar_chart(df: pd.DataFrame, x: str, y: str, title: str, xlabel: str, ylabel: str, path: Path, horizontal: bool = False):
    fig, ax = plt.subplots(figsize=(8.5, 4.6))
    if horizontal:
        ax.barh(df[y], df[x])
        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
    else:
        ax.bar(df[x], df[y])
        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
    ax.set_title(title)
    ax.grid(True, alpha=0.2, axis="y")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def _save_pie_chart(labels: list[str], values: list[float], title: str, path: Path):
    fig, ax = plt.subplots(figsize=(7.0, 5.0))
    ax.pie(values, labels=labels, autopct="%1.1f%%", startangle=90)
    ax.set_title(title)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def create_dashboard_charts(results: dict, chart_dir: Path) -> list[tuple[str, Path, str]]:
    chart_dir.mkdir(parents=True, exist_ok=True)
    base = results["base"]
    charts: list[tuple[str, Path, str]] = []

    df_growth = pd.DataFrame(base.get("growth_curve", []))
    df_feed = pd.DataFrame(base.get("feeding_plan", []))
    df_cost = pd.DataFrame(base.get("cost_curve", []))
    schedule = base.get("production_schedule", {})

    if not df_growth.empty and {"day", "weight_g"}.issubset(df_growth.columns):
        p = chart_dir / "crescimento.png"
        _save_line_chart(df_growth, "day", "weight_g", "Crescimento dos peixes ao longo do ciclo", "Peso (g)", p)
        charts.append(("Crescimento ao longo do ciclo", p, "Curva de evolução do peso médio dos peixes."))

    if not df_growth.empty and {"day", "biomass_kg"}.issubset(df_growth.columns):
        p = chart_dir / "biomassa.png"
        _save_line_chart(df_growth, "day", "biomass_kg", "Evolução da biomassa no ciclo", "Biomassa (kg)", p)
        charts.append(("Biomassa ao longo do ciclo", p, "Curva de biomassa acumulada ao longo do ciclo produtivo."))

    if not df_feed.empty and {"phase_name", "feed_phase_kg"}.issubset(df_feed.columns):
        p = chart_dir / "racao_fase.png"
        _save_bar_chart(df_feed, "feed_phase_kg", "phase_name", "Consumo de ração por fase", "Ração (kg)", "Fase", p, horizontal=True)
        charts.append(("Consumo de ração por fase", p, "Distribuição do consumo de ração entre as fases do cultivo."))

    if not df_feed.empty and {"phase_name", "feed_phase_cost"}.issubset(df_feed.columns):
        p = chart_dir / "custo_fase.png"
        _save_bar_chart(df_feed, "phase_name", "feed_phase_cost", "Custo de ração por fase", "Fase", "Custo (R$)", p)
        charts.append(("Custo de ração por fase", p, "Comparativo do custo de ração por fase de cultivo."))

    if not df_cost.empty and {"day", "cumulative_feed_cost"}.issubset(df_cost.columns):
        p = chart_dir / "custo_acumulado.png"
        _save_line_chart(df_cost, "day", "cumulative_feed_cost", "Custo acumulado da ração", "Custo acumulado (R$)", p)
        charts.append(("Custo acumulado da ração", p, "Evolução do custo acumulado da ração ao longo do ciclo."))

    labels = ["Ração", "Alevinos", "Energia adicional", "Mão de obra", "Outros custos", "Aeração"]
    values = [
        float(base.get("feed_cost_cycle", 0.0)),
        float(base.get("fingerling_cost_cycle", 0.0)),
        float(base.get("electricity_cost_cycle_effective", 0.0)),
        float(base.get("labor_cost_cycle_effective", 0.0)),
        float(base.get("other_costs_cycle_effective", 0.0)),
        float(base.get("aeration_energy_cost_cycle", 0.0)),
    ]
    filtered = [(l, v) for l, v in zip(labels, values) if v > 0]
    if filtered:
        p = chart_dir / "opex.png"
        _save_pie_chart([x[0] for x in filtered], [x[1] for x in filtered], "Composição do OPEX por ciclo", p)
        charts.append(("Composição do OPEX", p, "Participação relativa dos principais componentes do custo operacional por ciclo."))

    if schedule.get("production_strategy") == "Escalonada":
        harvests_per_year = int(float(schedule.get("harvests_per_year", 0) or 0))
        if harvests_per_year > 0:
            months = ["Jan", "Fev", "Mar", "Abr", "Mai", "Jun", "Jul", "Ago", "Set", "Out", "Nov", "Dez"]
            values = [0.0] * 12
            for i in range(min(harvests_per_year, 12)):
                values[i] = float(schedule.get("production_per_harvest_kg", 0.0))
            df_month = pd.DataFrame({"mês": months, "produção_kg": values})
            p = chart_dir / "producao_mensal.png"
            _save_bar_chart(df_month, "mês", "produção_kg", "Distribuição da produção mensal", "Mês", "Produção (kg)", p)
            charts.append(("Produção mensal estimada", p, "Distribuição mensal estimada de produção no modo escalonado."))

    return charts


def _add_feed_table(doc: Document, feeding_plan: list[dict]):
    doc.add_paragraph(style="Heading 2").add_run("Tabela do plano alimentar por fase")
    headers = [
        "Fase", "Faixa de peso", "Dias", "PB (%)", "Pellet (mm)", "Preço/kg", "FCR", "Ração fase (kg)", "Custo fase"
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header

    for row in feeding_plan:
        cells = table.add_row().cells
        cells[0].text = str(row.get("phase_name", "-"))
        cells[1].text = str(row.get("weight_range", "-"))
        cells[2].text = num(row.get("days_in_phase"), 0)
        cells[3].text = num(row.get("protein_percent"), 1)
        cells[4].text = num(row.get("pellet_mm"), 1)
        cells[5].text = brl(row.get("feed_price_per_kg"))
        cells[6].text = num(row.get("phase_fcr"), 2)
        cells[7].text = num(row.get("feed_phase_kg"))
        cells[8].text = brl(row.get("feed_phase_cost"))


def export_markdown_to_docx(input_md: Path, output_docx: Path, title: str, author: str) -> Path:
    doc = Document()
    configure_page(doc)
    add_cover(doc, title, author)
    add_toc(doc)
    doc.add_paragraph("Esta exportação simples foi substituída pelo relatório premium.")
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return output_docx


def export_dashboard_report_to_docx(
    results: dict,
    output_docx: Path,
    title: str,
    author: str,
    profile: str = "Produtor",
) -> Path:
    base = results["base"]
    inp = base["input"]
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    chart_dir = output_docx.parent / "_chart_cache"

    doc = Document()
    configure_page(doc)
    add_cover(doc, title, author)
    add_toc(doc)

    doc.add_paragraph(style="Heading 1").add_run("Resumo executivo")
    add_bullets(doc, [
        f"Perfil do relatório: {profile}",
        f"Produção por ciclo: {num(base.get('production_per_cycle_kg'))} kg",
        f"Produção anual: {num(base.get('production_per_year_kg'))} kg",
        f"Receita por ciclo: {brl(base.get('revenue_cycle'))}",
        f"OPEX por ciclo: {brl(base.get('opex_cycle'))}",
        f"Margem por ciclo: {brl(base.get('gross_margin_cycle'))}",
        f"Custo por kg: {brl(base.get('cost_per_kg'))}/kg",
        f"Payback estimado: {num(base.get('payback_years'), 2)} anos",
        f"CAPEX considerado: {brl(base.get('capex_total_effective'))}",
    ])

    profile_sections = build_profile_sections(results, profile)

    doc.add_paragraph(style="Heading 1").add_run("Leitura específica para o perfil")
    add_bullets(doc, [f"Perfil selecionado: {profile}", f"Enfoque do relatório: {profile_sections.get('title', '-')}"])
    add_bullets(doc, list(profile_sections.get("highlights", [])))

    doc.add_paragraph(style="Heading 1").add_run("Recomendações específicas para o perfil")
    add_bullets(doc, list(profile_sections.get("recommendations", [])))

    doc.add_paragraph(style="Heading 1").add_run("Pontos de decisão e próximos passos")
    add_bullets(doc, list(profile_sections.get("decisions", [])))

    doc.add_paragraph(style="Heading 1").add_run("Dimensionamento do sistema")
    add_bullets(doc, [
        f"Espécie: {inp.get('species', '-')}",
        f"Sistema: {inp.get('system_type', '-')}",
        f"Número de unidades: {inp.get('number_of_units', '-')}",
        f"Volume total: {num(base.get('total_volume_m3'))} m³",
        f"Biomassa final: {num(base.get('final_biomass_total_kg'))} kg",
        f"Peixes colhidos: {num(base.get('fish_harvested'))}",
        f"Alevinos para compra/ciclo: {num(base.get('fingerlings_purchase_cycle'), 0)}",
        f"Dias do ciclo: {num(base.get('estimated_cycle_days'), 0)}",
        f"Ganho diário ajustado: {num(base.get('adjusted_daily_growth_g'))} g/dia",
    ])

    schedule = base.get("production_schedule", {})
    doc.add_paragraph(style="Heading 1").add_run("Estratégia de produção")
    if schedule.get("production_strategy") == "Escalonada":
        add_bullets(doc, [
            f"Modo operacional: Escalonada",
            f"Intervalo entre despescas: {num(schedule.get('harvest_interval_months'), 2)} mês(es)",
            f"Lotes em paralelo: {num(schedule.get('batches_in_parallel'), 0)}",
            f"Tanques por lote: {num(schedule.get('units_per_batch_exact'), 2)}",
            f"Produção por despesca: {num(schedule.get('production_per_harvest_kg'))} kg",
            f"Receita por despesca: {brl(schedule.get('revenue_per_harvest'))}",
            f"Alevinos arredondados por despesca: {num(schedule.get('fingerlings_purchase_per_harvest'), 0)}",
        ])
    else:
        add_bullets(doc, [
            "Modo operacional: Ciclos simultâneos",
            f"Colheitas por ano: {num(schedule.get('harvests_per_year'), 2)}",
            f"Produção por colheita: {num(schedule.get('production_per_harvest_kg'))} kg",
        ])

    doc.add_paragraph(style="Heading 1").add_run("Dashboard visual e gráficos")
    charts = create_dashboard_charts(results, chart_dir)
    for chart_title, chart_path, caption in charts:
        add_chart_image(doc, chart_title, chart_path, caption)

    doc.add_paragraph(style="Heading 1").add_run("Plano alimentar por fase")
    _add_feed_table(doc, base.get("feeding_plan", []))

    doc.add_paragraph(style="Heading 1").add_run("Custos e indicadores econômicos")
    add_bullets(doc, [
        f"Custo de ração por ciclo: {brl(base.get('feed_cost_cycle'))}",
        f"Custo de ração anual: {brl(base.get('feed_cost_year'))}",
        f"Custo de alevinos por ciclo: {brl(base.get('fingerling_cost_cycle'))}",
        f"Custo de aeração anual: {brl(base.get('aeration_energy_cost_year'))}",
        f"Energia adicional anual considerada: {brl(base.get('electricity_cost_year'))}",
        f"Mão de obra anual considerada: {brl(base.get('labor_cost_year'))}",
        f"Outros custos anuais considerados: {brl(base.get('other_costs_year'))}",
        f"OPEX anual: {brl(base.get('opex_year'))}",
        f"FCR implícito do ciclo: {num(base.get('implied_fcr'))}",
        f"CAPEX fixo considerado: {brl(base.get('capex_fixed_total_effective'))}",
        f"CAPEX variável considerado: {brl(base.get('capex_variable_total_effective'))}",
    ])

    doc.add_paragraph(style="Heading 1").add_run("Observações finais")
    doc.add_paragraph(str(inp.get("notes", "Nenhuma observação adicional informada.")))

    doc.save(str(output_docx))
    return output_docx


def export_docx_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    cmd = ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(result.stderr or result.stdout or "Falha ao converter PDF")
    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.exists():
        raise RuntimeError("PDF não gerado.")
    return pdf_path
