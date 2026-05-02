from __future__ import annotations

import base64
import math
import os
from dataclasses import fields
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
import streamlit as st

def get_image_base64(image_path: str) -> str:
    path = Path(__file__).resolve().parent / image_path
    if not path.exists():
        return ""
    return base64.b64encode(path.read_bytes()).decode()


def _slugify_filename(value: str, fallback: str = "projeto_tilapia_dashboard") -> str:
    """Gera um nome seguro para arquivos a partir do nome do projeto."""
    raw = str(value or "").strip()
    if not raw:
        raw = fallback
    replacements = {
        "á": "a", "à": "a", "â": "a", "ã": "a", "ä": "a",
        "é": "e", "ê": "e", "è": "e", "ë": "e",
        "í": "i", "ì": "i", "î": "i", "ï": "i",
        "ó": "o", "ô": "o", "õ": "o", "ò": "o", "ö": "o",
        "ú": "u", "ù": "u", "û": "u", "ü": "u",
        "ç": "c",
        "Á": "A", "À": "A", "Â": "A", "Ã": "A", "Ä": "A",
        "É": "E", "Ê": "E", "È": "E", "Ë": "E",
        "Í": "I", "Ì": "I", "Î": "I", "Ï": "I",
        "Ó": "O", "Ô": "O", "Õ": "O", "Ò": "O", "Ö": "O",
        "Ú": "U", "Ù": "U", "Û": "U", "Ü": "U",
        "Ç": "C",
    }
    for src, dst in replacements.items():
        raw = raw.replace(src, dst)
    safe = []
    for ch in raw:
        if ch.isalnum():
            safe.append(ch.lower())
        elif ch in (" ", "-", "_"):
            safe.append("_")
    name = "".join(safe)
    while "__" in name:
        name = name.replace("__", "_")
    name = name.strip("_")
    return name or fallback


from aqua_project_agent_gui_dashboard_v1.calculator import (
    DEFAULT_FEEDING_CURVE,
    FOUNTAIN_LIBRARY,
    LOBULAR_BLOWER_LIBRARY,
    PADDLEWHEEL_LIBRARY,
    RADIAL_BLOWER_LIBRARY,
    altitude_transfer_factor,
    calculate_dashboard_project,
    recommended_depth_for_structure,
    surface_aeration_base_factor,
)
from aqua_project_agent_gui_dashboard_v1.exporter import (
    export_docx_to_pdf,
    export_markdown_to_docx,
)
from aqua_project_agent_gui_dashboard_v1.models import DashboardProjectInput
from aqua_project_agent_gui_dashboard_v1.report import (
    build_professional_project_report,
    brl,
    num,
)

from auth.supabase_auth import is_supabase_configured, sign_in_with_password
from data_access.local_store import (
    build_project_payload,
    delete_project,
    duplicate_project,
    list_projects,
    load_project,
    save_project,
)
from data_access.supabase_store import (
    delete_project_remote,
    duplicate_project_remote,
    list_projects_remote,
    load_project_remote,
    save_project_remote,
)

st.set_page_config(
    page_title="Aqua Project Agent Dashboard",
    layout="wide",
    initial_sidebar_state="expanded",
)

BG_IMAGE_BASE64 = get_image_base64("assets/tilapia_scale_bg.png")

if BG_IMAGE_BASE64:
    app_background_css = f"""
    background-image:
        linear-gradient(
            rgba(238, 247, 244, 0.22),
            rgba(238, 247, 244, 0.38)
        ),
        url("data:image/png;base64,{BG_IMAGE_BASE64}") !important;
    background-size: cover !important;
    background-position: center center !important;
    background-repeat: no-repeat !important;
    background-attachment: fixed !important;
    """
else:
    app_background_css = """
    background: linear-gradient(180deg, #eef7f5 0%, #dcefeb 100%) !important;
    """

st.markdown(
    f"""
<style>
/* ============================================================
   FUNDO PRINCIPAL COM TEXTURA DE ESCAMAS DA TILÁPIA
   ============================================================ */

.stApp {{
    {app_background_css}
}}

[data-testid="stAppViewContainer"],
[data-testid="stMain"] {{
    background: transparent !important;
}}

[data-testid="stHeader"] {{
    background: rgba(238, 247, 244, 0.58) !important;
    backdrop-filter: blur(6px) !important;
}}

.block-container {{
    background: rgba(238, 247, 244, 0.52) !important;
    border-radius: 18px !important;
    padding-top: 1.2rem !important;
    padding-bottom: 2rem !important;
    max-width: 1280px !important;
}}

/* ============================================================
   TIPOGRAFIA E CONTRASTE GERAL
   ============================================================ */

main,
main p,
main span,
main div,
main li {{
    color: #102A43 !important;
}}

h1, h2, h3, h4, h5, h6 {{
    color: #0F2D3A !important;
    font-weight: 800 !important;
}}

label,
[data-testid="stWidgetLabel"],
[data-testid="stWidgetLabel"] p,
[data-testid="stWidgetLabel"] span,
[data-testid="stMarkdownContainer"] p,
[data-testid="stMarkdownContainer"] span,
[data-testid="stMarkdownContainer"] li {{
    color: #102A43 !important;
}}

[data-testid="stWidgetLabel"] {{
    color: #334E5C !important;
    font-weight: 800 !important;
}}

/* ============================================================
   CAMPOS DE FORMULÁRIO — CINZA-AZULADO + CURSOR VERMELHO
   Regras principais:
   - sem foco: fundo cinza-azulado claro
   - em edição/foco: fundo branco, borda verde
   - cursor/caret de texto: vermelho
   - texto digitado: azul escuro
   ============================================================ */

input,
textarea,
[data-baseweb="input"] > div,
[data-baseweb="textarea"] > div,
div[data-baseweb="select"] > div {{
    background-color: #F3F7FA !important;
    color: #0F2D3A !important;
    border: 1px solid #8FA6B3 !important;
    caret-color: #D62828 !important;
}}

input:hover,
textarea:hover,
[data-baseweb="input"] > div:hover,
[data-baseweb="textarea"] > div:hover,
div[data-baseweb="select"] > div:hover {{
    background-color: #EAF2F6 !important;
    border-color: #5F8796 !important;
}}

/* Texto interno dos campos */
[data-baseweb="input"] input,
[data-baseweb="textarea"] textarea,
[data-testid="stNumberInput"] input,
div[data-baseweb="select"] input {{
    background-color: #F3F7FA !important;
    color: #0F2D3A !important;
    caret-color: #D62828 !important;
}}

/* Cursor do mouse sobre campos editáveis */
[data-baseweb="input"],
[data-baseweb="input"] *,
[data-baseweb="textarea"],
[data-baseweb="textarea"] *,
input,
textarea {{
    cursor: text !important;
}}

div[data-baseweb="select"],
div[data-baseweb="select"] * {{
    cursor: pointer !important;
}}

/* Placeholders */
input::placeholder,
textarea::placeholder {{
    color: #60717D !important;
    opacity: 1 !important;
}}

/* Selectbox */
div[data-baseweb="select"] span,
div[data-baseweb="select"] input,
div[data-baseweb="select"] div {{
    color: #0F2D3A !important;
}}

div[data-baseweb="select"] svg {{
    color: #0F2D3A !important;
    fill: #0F2D3A !important;
}}

/* Menu aberto dos selectboxes */
div[data-baseweb="popover"],
div[data-baseweb="popover"] div,
ul[role="listbox"],
li[role="option"],
li[role="option"] div,
li[role="option"] span {{
    background-color: #FFFFFF !important;
    color: #0F2D3A !important;
}}

li[role="option"]:hover,
li[role="option"]:hover div,
li[role="option"]:hover span {{
    background-color: #DFF3EE !important;
    color: #0F2D3A !important;
}}

/* Campo desabilitado, como Espécie */
input:disabled,
textarea:disabled {{
    background-color: #E6EEF2 !important;
    color: #60717D !important;
    opacity: 1 !important;
}}

/* Foco/edição: campo branco com borda verde */
input:focus,
textarea:focus,
[data-baseweb="input"]:focus-within > div,
[data-baseweb="textarea"]:focus-within > div,
div[data-baseweb="select"]:focus-within > div {{
    background-color: #FFFFFF !important;
    color: #0F2D3A !important;
    border: 2px solid #0F766E !important;
    box-shadow: 0 0 0 2px rgba(15, 118, 110, 0.18) !important;
    caret-color: #D62828 !important;
}}

[data-baseweb="input"]:focus-within input,
[data-baseweb="textarea"]:focus-within textarea,
[data-testid="stNumberInput"]:focus-within input {{
    background-color: #FFFFFF !important;
    color: #0F2D3A !important;
    caret-color: #D62828 !important;
}}

/* Number input: botões + e - */
[data-testid="stNumberInput"] button {{
    background-color: #0F2D3A !important;
    color: #FFFFFF !important;
    border-color: #0F2D3A !important;
    cursor: pointer !important;
}}

[data-testid="stNumberInput"] button svg {{
    color: #FFFFFF !important;
    fill: #FFFFFF !important;
}}

/* ============================================================
   ABAS
   ============================================================ */

button[data-baseweb="tab"] p,
button[data-baseweb="tab"] span {{
    color: #4E6472 !important;
    font-weight: 700 !important;
}}

button[data-baseweb="tab"][aria-selected="true"] p,
button[data-baseweb="tab"][aria-selected="true"] span {{
    color: #0F766E !important;
    font-weight: 800 !important;
}}

/* ============================================================
   SIDEBAR
   ============================================================ */

[data-testid="stSidebar"] {{
    background: linear-gradient(180deg, #075E66 0%, #0F766E 100%) !important;
}}

[data-testid="stSidebar"] label,
[data-testid="stSidebar"] .stMarkdown,
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 {{
    color: #FFFFFF !important;
}}

[data-testid="stSidebar"] input,
[data-testid="stSidebar"] textarea,
[data-testid="stSidebar"] [data-baseweb="input"] > div,
[data-testid="stSidebar"] [data-baseweb="textarea"] > div,
[data-testid="stSidebar"] div[data-baseweb="select"] > div {{
    background-color: #F3F7FA !important;
    color: #0F2D3A !important;
    caret-color: #D62828 !important;
}}

[data-testid="stSidebar"] div[data-baseweb="select"] span,
[data-testid="stSidebar"] div[data-baseweb="select"] input {{
    color: #0F2D3A !important;
}}

/* ============================================================
   BOTÕES GERAIS
   ============================================================ */

.stButton > button,
button[kind="secondary"],
button[kind="primary"] {{
    color: #FFFFFF !important;
    background-color: #0F2D3A !important;
    border: 1px solid #0F2D3A !important;
    font-weight: 700 !important;
    border-radius: 10px !important;
}}

.stButton > button p,
.stButton > button span,
button[kind="secondary"] p,
button[kind="secondary"] span,
button[kind="primary"] p,
button[kind="primary"] span {{
    color: #FFFFFF !important;
}}

.stButton > button:hover,
button[kind="secondary"]:hover,
button[kind="primary"]:hover {{
    background-color: #0F766E !important;
    border-color: #0F766E !important;
    color: #FFFFFF !important;
}}

.stButton > button:disabled,
button:disabled {{
    color: #CBD5E1 !important;
    background-color: #1F2937 !important;
    border: 1px solid #334155 !important;
    opacity: 1 !important;
}}

.stButton > button:disabled p,
.stButton > button:disabled span,
button:disabled p,
button:disabled span {{
    color: #CBD5E1 !important;
    opacity: 1 !important;
}}

/* ============================================================
   CARDS, KPIS E SEÇÕES
   ============================================================ */

.kpi-card {{
    background: rgba(255, 255, 255, 0.90) !important;
    padding: 18px !important;
    border-radius: 18px !important;
    box-shadow: 0 6px 18px rgba(15, 45, 58, 0.12) !important;
    border-left: 6px solid #0F766E !important;
    margin-bottom: 12px !important;
}}

.kpi-title {{
    font-size: 14px !important;
    color: #4E6472 !important;
    margin-bottom: 6px !important;
    font-weight: 700 !important;
}}

.kpi-value {{
    font-size: 28px !important;
    font-weight: 800 !important;
    color: #0F2D3A !important;
}}

.section-box {{
    background: rgba(255, 255, 255, 0.86) !important;
    border: 1px solid #C9D8D4 !important;
    border-radius: 18px !important;
    padding: 18px !important;
    box-shadow: 0 6px 18px rgba(15, 45, 58, 0.10) !important;
    margin-bottom: 18px !important;
}}

.hero-box {{
    background: linear-gradient(135deg, rgba(15,76,92,0.96), rgba(15,118,110,0.94)) !important;
    color: white !important;
    border-radius: 22px !important;
    padding: 28px 30px !important;
    box-shadow: 0 8px 24px rgba(0,0,0,0.16) !important;
    margin-bottom: 18px !important;
}}

.hero-title {{
    font-size: 34px !important;
    font-weight: 800 !important;
    margin-bottom: 6px !important;
    color: #FFFFFF !important;
}}

.hero-subtitle {{
    font-size: 18px !important;
    opacity: 0.95 !important;
    margin-bottom: 0 !important;
    color: #FFFFFF !important;
}}

/* Métricas nativas */
[data-testid="stMetric"],
[data-testid="stMetric"] div,
[data-testid="stMetric"] label,
[data-testid="stMetric"] span {{
    color: #102A43 !important;
}}

[data-testid="stMetricValue"],
[data-testid="stMetricValue"] div {{
    color: #0F2D3A !important;
    font-weight: 800 !important;
}}

/* Alertas */
[data-testid="stAlert"] {{
    background-color: rgba(217, 238, 247, 0.94) !important;
    color: #102A43 !important;
    border-radius: 12px !important;
}}

[data-testid="stAlert"] p,
[data-testid="stAlert"] span,
[data-testid="stAlert"] div {{
    color: #102A43 !important;
}}


/* ============================================================
   BLOCO DE USUÁRIO CONECTADO
   ============================================================ */

.connected-user-box {{
    background: rgba(255, 255, 255, 0.18) !important;
    border: 1px solid rgba(255, 255, 255, 0.34) !important;
    border-radius: 12px !important;
    padding: 10px 12px !important;
    margin-top: 8px !important;
    margin-bottom: 8px !important;
}}

.connected-user-label {{
    color: #E6FFFA !important;
    font-size: 0.82rem !important;
    font-weight: 800 !important;
    margin-bottom: 2px !important;
}}

.connected-user-email {{
    color: #FFFFFF !important;
    font-size: 0.92rem !important;
    font-weight: 800 !important;
    word-break: break-word !important;
}}


/* ============================================================
   BARRAS DE ROLAGEM — ALTO CONTRASTE
   ============================================================ */
html, body,
[data-testid="stAppViewContainer"],
[data-testid="stMain"],
[data-testid="stSidebar"],
[data-testid="stSidebarContent"],
section[data-testid="stSidebar"] {{
    scrollbar-color: #011F2A #063E43 !important;
    scrollbar-width: auto !important;
}}

html::-webkit-scrollbar,
body::-webkit-scrollbar,
[data-testid="stAppViewContainer"]::-webkit-scrollbar,
[data-testid="stMain"]::-webkit-scrollbar,
*::-webkit-scrollbar {{
    width: 20px !important;
    height: 20px !important;
}}

html::-webkit-scrollbar-track,
body::-webkit-scrollbar-track,
[data-testid="stAppViewContainer"]::-webkit-scrollbar-track,
[data-testid="stMain"]::-webkit-scrollbar-track,
*::-webkit-scrollbar-track {{
    background: #063E43 !important;
    border-radius: 12px !important;
    box-shadow: inset 0 0 0 1px rgba(255,255,255,0.28) !important;
}}

html::-webkit-scrollbar-thumb,
body::-webkit-scrollbar-thumb,
[data-testid="stAppViewContainer"]::-webkit-scrollbar-thumb,
[data-testid="stMain"]::-webkit-scrollbar-thumb,
*::-webkit-scrollbar-thumb {{
    background: #011F2A !important;
    border-radius: 12px !important;
    border: 3px solid #19A39A !important;
    min-height: 56px !important;
}}

html::-webkit-scrollbar-thumb:hover,
body::-webkit-scrollbar-thumb:hover,
[data-testid="stAppViewContainer"]::-webkit-scrollbar-thumb:hover,
[data-testid="stMain"]::-webkit-scrollbar-thumb:hover,
*::-webkit-scrollbar-thumb:hover {{
    background: #000F16 !important;
    border-color: #7AE0DA !important;
}}

html::-webkit-scrollbar-corner,
body::-webkit-scrollbar-corner,
[data-testid="stAppViewContainer"]::-webkit-scrollbar-corner,
[data-testid="stMain"]::-webkit-scrollbar-corner,
*::-webkit-scrollbar-corner {{
    background: #063E43 !important;
}}

/* Sidebar: mantém o mesmo padrão forte da barra principal */
[data-testid="stSidebar"] *::-webkit-scrollbar,
[data-testid="stSidebar"]::-webkit-scrollbar,
section[data-testid="stSidebar"] *::-webkit-scrollbar {{
    width: 20px !important;
}}

[data-testid="stSidebar"] *::-webkit-scrollbar-track,
[data-testid="stSidebar"]::-webkit-scrollbar-track,
section[data-testid="stSidebar"] *::-webkit-scrollbar-track {{
    background: #063E43 !important;
    box-shadow: inset 0 0 0 1px rgba(255,255,255,0.28) !important;
}}

[data-testid="stSidebar"] *::-webkit-scrollbar-thumb,
[data-testid="stSidebar"]::-webkit-scrollbar-thumb,
section[data-testid="stSidebar"] *::-webkit-scrollbar-thumb {{
    background: #011F2A !important;
    border: 3px solid #19A39A !important;
}}

[data-testid="stSidebar"] *::-webkit-scrollbar-thumb:hover,
[data-testid="stSidebar"]::-webkit-scrollbar-thumb:hover,
section[data-testid="stSidebar"] *::-webkit-scrollbar-thumb:hover {{
    background: #000F16 !important;
    border-color: #7AE0DA !important;
}}

/* ============================================================
   RESPONSIVO
   ============================================================ */

@media screen and (max-width: 768px) {{
    .block-container {{
        padding-left: 1rem !important;
        padding-right: 1rem !important;
        background: rgba(238, 247, 244, 0.78) !important;
    }}

    main h1 {{
        font-size: 2.1rem !important;
        line-height: 1.15 !important;
    }}

    main h2 {{
        font-size: 1.7rem !important;
        line-height: 1.2 !important;
    }}

    main h3 {{
        font-size: 1.35rem !important;
        line-height: 1.25 !important;
    }}

    label,
    [data-testid="stWidgetLabel"],
    [data-testid="stWidgetLabel"] p,
    [data-testid="stWidgetLabel"] span {{
        color: #334E5C !important;
        font-size: 0.95rem !important;
        font-weight: 800 !important;
    }}

    input,
    textarea,
    div[data-baseweb="select"] > div {{
        min-height: 48px !important;
        font-size: 1rem !important;
    }}

    .stButton > button {{
        min-height: 48px !important;
        font-size: 1rem !important;
    }}
}}
</style>
""",
    unsafe_allow_html=True,
)

def _ui_geometry(system_type: str, unit_volume_m3: float, length_m: float, width_m: float, depth_m: float) -> dict:
    depth = max(float(depth_m), 0.2)
    if system_type == "Circular revestido":
        volume = max(float(unit_volume_m3), 0.1)
        diameter = (4.0 * volume / (3.141592653589793 * depth)) ** 0.5
        area = 3.141592653589793 * (diameter ** 2) / 4.0
        return {"volume_m3": volume, "diameter_m": diameter, "area_m2": area, "length_m": diameter, "width_m": diameter}
    length = max(float(length_m), 0.1)
    width = max(float(width_m), 0.1)
    area = length * width
    volume = area * depth
    return {"volume_m3": volume, "diameter_m": 0.0, "area_m2": area, "length_m": length, "width_m": width}


def _model_options(library: list[dict], key: str = "model") -> list[str]:
    return [item[key] for item in library]



def _safe_float_for_ui(value, default: float | None = None) -> float | None:
    """Converte valores numéricos em float, aceitando formatos brasileiros."""
    try:
        if value is None or value == "":
            return default
        if isinstance(value, str):
            value = value.replace("R$", "").replace("%", "").replace("x", "").strip()
            value = value.replace(".", "").replace(",", ".")
        return float(value)
    except Exception:
        return default


def _first_existing(mapping: dict, keys: list[str], default=None):
    """Retorna o primeiro valor não vazio entre várias chaves possíveis."""
    if not isinstance(mapping, dict):
        return default
    for key in keys:
        value = mapping.get(key)
        if value not in (None, ""):
            return value
    return default


def _parse_weight_range_g(row: dict) -> tuple[float | None, float | None]:
    """Extrai faixa de peso em gramas a partir do feeding_plan."""
    min_g = _safe_float_for_ui(_first_existing(row, ["min_g", "phase_min_g", "Peso mín. (g)"]), None)
    max_g = _safe_float_for_ui(_first_existing(row, ["max_g", "phase_max_g", "Peso máx. (g)"]), None)
    if min_g is not None and max_g is not None:
        return min_g, max_g

    text = str(row.get("weight_range", row.get("Faixa de peso", "")))
    import re
    nums = re.findall(r"\d+(?:[\.,]\d+)?", text)
    if len(nums) >= 2:
        return _safe_float_for_ui(nums[0], None), _safe_float_for_ui(nums[1], None)
    return None, None


def _format_pct_for_table(value) -> str:
    number = _safe_float_for_ui(value, None)
    if number is None:
        return "-"
    return f"{num(number, 1)}%"


def _format_number_for_table(value, decimals: int = 2) -> str:
    number = _safe_float_for_ui(value, None)
    if number is None:
        return "-"
    return num(number, decimals)


def _format_money_for_table(value) -> str:
    number = _safe_float_for_ui(value, None)
    if number is None:
        return "-"
    return brl(number)


def _build_aeration_phase_table_rows(base_data: dict, form_data: dict) -> list[dict]:
    """
    Consolida a operação da aeração por fase e preenche colunas que o cálculo
    original pode não retornar diretamente: biomassa média, demanda de O₂ e dias.
    """
    if not isinstance(base_data, dict):
        return []

    aer_rows = base_data.get("aeration_phase_operation", []) or []
    feeding_plan = base_data.get("feeding_plan", []) or []
    if not isinstance(aer_rows, list):
        aer_rows = []
    if not isinstance(feeding_plan, list):
        feeding_plan = []

    number_of_units = _safe_float_for_ui(form_data.get("number_of_units"), 1.0) or 1.0
    survival = _safe_float_for_ui(form_data.get("survival_rate"), 0.90) or 0.90
    survival = survival if survival > 0 else 0.90
    target_weight_kg = (_safe_float_for_ui(form_data.get("target_weight_g"), 1000.0) or 1000.0) / 1000.0
    target_weight_kg = target_weight_kg if target_weight_kg > 0 else 1.0

    total_harvest_biomass = _safe_float_for_ui(base_data.get("production_per_cycle_kg"), None)
    if total_harvest_biomass is None:
        final_per_tank = _safe_float_for_ui(base_data.get("final_biomass_per_tank_kg"), 0.0) or 0.0
        total_harvest_biomass = final_per_tank * number_of_units

    initial_fish_count = total_harvest_biomass / target_weight_kg / survival if total_harvest_biomass else 0.0
    oxygen_base_mg = _safe_float_for_ui(form_data.get("oxygen_demand_mg_per_kg_h"), 550.0) or 550.0
    safety_factor = 1.0 + ((_safe_float_for_ui(form_data.get("aeration_safety_factor_pct"), 0.0) or 0.0) / 100.0)
    hours_per_day = _safe_float_for_ui(form_data.get("aeration_hours_per_day"), 24.0) or 24.0
    energy_price = _safe_float_for_ui(form_data.get("electricity_price_kwh"), 0.45) or 0.45
    default_equipment = _safe_float_for_ui(base_data.get("required_aerators"), None)

    n_rows = max(len(aer_rows), len(feeding_plan))
    output: list[dict] = []
    for idx in range(n_rows):
        aer = aer_rows[idx] if idx < len(aer_rows) and isinstance(aer_rows[idx], dict) else {}
        feed = feeding_plan[idx] if idx < len(feeding_plan) and isinstance(feeding_plan[idx], dict) else {}

        phase = _first_existing(aer, ["Fase", "phase_name", "fase"], None) or feed.get("phase_name") or f"Fase {idx + 1}"

        biomass = _safe_float_for_ui(_first_existing(aer, ["Biomassa média (kg)", "biomass_mean_kg", "average_biomass_kg", "Biomassa (kg)"]), None)
        if biomass is None:
            min_g, max_g = _parse_weight_range_g(feed)
            if min_g is not None and max_g is not None and initial_fish_count > 0:
                biomass = initial_fish_count * ((min_g + max_g) / 2.0) / 1000.0

        demand = _safe_float_for_ui(_first_existing(aer, ["Demanda O₂ (kg/h)", "Demanda O2 (kg/h)", "oxygen_demand_kg_h", "demand_kg_h", "phase_oxygen_demand_kg_h"]), None)
        if demand is None and biomass is not None:
            demand = biomass * oxygen_base_mg * safety_factor / 1_000_000.0

        modulation = _first_existing(aer, ["% de uso da capacidade", "% de modulação sugerida", "Modulação sugerida", "usage_pct", "capacity_usage_pct"], None)
        equipment = _safe_float_for_ui(_first_existing(aer, ["Equipamentos ativos", "Equipamentos operando", "active_equipment", "equipment_active"], default_equipment), None)
        power = _safe_float_for_ui(_first_existing(aer, ["Potência ativa da fase (kW)", "Potência média ativa da fase (kW)", "active_power_kw", "average_active_power_kw"], None), None)
        days = _safe_float_for_ui(_first_existing(aer, ["Dias", "days", "days_in_phase"], feed.get("days_in_phase")), None)
        cost = _safe_float_for_ui(_first_existing(aer, ["Custo da fase", "Custo da fase (R$)", "phase_cost", "aeration_cost_phase"], None), None)
        if cost is None and power is not None and days is not None:
            cost = power * hours_per_day * days * energy_price

        output.append({
            "Fase": phase,
            "Biomassa média (kg)": _format_number_for_table(biomass, 2),
            "Demanda O₂ (kg/h)": _format_number_for_table(demand, 3),
            "Modulação sugerida": _format_pct_for_table(modulation),
            "Equipamentos operando": _format_number_for_table(equipment, 0),
            "Potência média ativa (kW)": _format_number_for_table(power, 2),
            "Dias": _format_number_for_table(days, 0),
            "Custo da fase": _format_money_for_table(cost),
        })
    return output


def _apply_altitude_correction_to_aeration_results(base_data: dict, form_data: dict) -> dict:
    """
    Retorna um resumo da correção por altitude sem recalcular novamente os valores.

    A correção efetiva agora é feita no motor de cálculo (calculator.py).
    Esta função serve apenas para exibição na interface, evitando dupla penalização.
    """
    if not isinstance(base_data, dict):
        return {}

    altitude_m = _safe_float_for_ui(form_data.get("site_altitude_m"), 0.0) or 0.0
    altitude_factor = _safe_float_for_ui(
        base_data.get("altitude_transfer_factor", base_data.get("altitude_factor")),
        altitude_transfer_factor(altitude_m),
    ) or altitude_transfer_factor(altitude_m)

    installed_supply_sea_level = _safe_float_for_ui(
        base_data.get("installed_oxygen_supply_sea_level_kg_h"),
        base_data.get("installed_oxygen_supply_kg_h"),
    ) or 0.0
    installed_supply_corrected = _safe_float_for_ui(base_data.get("installed_oxygen_supply_kg_h"), 0.0) or 0.0
    oxygen_demand_total = _safe_float_for_ui(base_data.get("oxygen_demand_kg_h"), 0.0) or 0.0
    required_aerators = max(0, int(_safe_float_for_ui(base_data.get("required_aerators"), 0) or 0))
    ratio = installed_supply_corrected / oxygen_demand_total if oxygen_demand_total > 0 else None

    return {
        "altitude_m": altitude_m,
        "altitude_factor": altitude_factor,
        "required_aerators_original": required_aerators,
        "required_aerators_corrected": required_aerators,
        "installed_supply_original": installed_supply_sea_level,
        "installed_supply_same_equipment": installed_supply_sea_level * altitude_factor,
        "installed_supply_corrected": installed_supply_corrected,
        "oxygen_demand_total": oxygen_demand_total,
        "oxygen_offer_demand_ratio": ratio,
    }


def _estimate_preliminary_capex_from_form(form_data: dict, geometry: dict) -> dict:
    """Estimativa preliminar de CAPEX por composição, sem substituir orçamento real."""
    units = max(1.0, _safe_float_for_ui(form_data.get("number_of_units"), 1.0) or 1.0)
    unit_volume = max(0.1, float(geometry.get("volume_m3", form_data.get("unit_volume_m3", 100.0)) or 100.0))
    total_volume = units * unit_volume
    # Valor preliminar parametrizado. Deve ser revisado por cotação atualizada.
    tank_structure = total_volume * 500.0
    hydraulics = tank_structure * 0.16
    aeration = tank_structure * 0.12
    electrical = tank_structure * 0.10
    backup = tank_structure * 0.08
    civil_install = tank_structure * 0.14
    subtotal = tank_structure + hydraulics + aeration + electrical + backup + civil_install
    contingency = subtotal * 0.15
    total = subtotal + contingency
    return {
        "Tanques/estrutura": tank_structure,
        "Hidráulica e drenagem": hydraulics,
        "Aeração/oxigenação": aeration,
        "Elétrica/comando": electrical,
        "Backup/gerador": backup,
        "Obras civis/instalação": civil_install,
        "Contingência técnica": contingency,
        "CAPEX preliminar total": total,
        "Premissa": "Estimativa preliminar por volume útil e composição percentual. Carece levantamento de preços atualizado.",
    }


sample_data = {
    "project_name": "Projeto tilápia tanques revestidos",
    "author_name": "Luiz Henrique Sousa Salgado",
    "region_focus": "Brasil, com ênfase no Sudeste",
    "system_type": "Circular revestido",
    "species": "Tilápia",
    "number_of_units": 9,
    "unit_volume_m3": 100.0,
    "water_depth_m": 1.2,
    "tank_length_m": 10.0,
    "tank_width_m": 10.0,
    "density_kg_m3": 30.0,
    "survival_rate": 0.90,
    "cycles_per_year": 2.0,
    "production_strategy": "Escalonada",
    "scheduling_basis": "Intervalo entre despescas",
    "harvest_interval_months": 1.0,
    "manual_parallel_batches": 9,
    "desired_units_per_batch": 1,
    "sale_price_per_kg": 10.0,
    "fingerling_price": 0.30,
    "fingerling_weight_kg": 1.0,
    "electricity_cost_cycle": 12000.0,
    "labor_cost_cycle": 18000.0,
    "other_costs_cycle": 15000.0,
    "capex_total": 450000.0,
    "capex_component_mode_note": "Estimativa preliminar. Atualizar com cotação real de tanques, hidráulica, elétrica, aeração, backup e instalação.",
    "cost_scaling_mode": "Fixos (não escalar)",
    "cost_reference_units": 9,
    "economic_model_mode": "Simplificado",
    "electricity_cost_fixed_cycle": 0.0,
    "electricity_cost_per_unit_cycle": 1500.0,
    "labor_cost_fixed_cycle": 0.0,
    "labor_cost_per_unit_cycle": 2250.0,
    "other_cost_fixed_cycle": 0.0,
    "other_cost_per_unit_cycle": 1875.0,
    "capex_fixed_total": 0.0,
    "capex_per_unit": 56250.0,
    "initial_weight_g": 1.0,
    "target_weight_g": 1000.0,
    "water_temperature_c": 28.0,
    "base_daily_growth_g": 4.5,
    "growth_curve_adjustment_pct": 100.0,
    "oxygen_demand_mg_per_kg_h": 550.0,
    "site_altitude_m": 0.0,
    "target_do_mg_l": 5.0,
    "field_efficiency_pct": 85.0,
    "aeration_safety_factor_pct": 20.0,
    "aeration_hours_per_day": 24.0,
    "electricity_price_kwh": 0.45,
    "aeration_mode": "Automático",
    "automatic_aeration_technology": "Chafariz",
    "blower_type": "Automático",
    "diffusion_efficiency_pct": 12.0,
    "aeration_power_mode": "Potência modulada por fase",
    "aeration_control_strategy": "Automático",
    "blower_min_operational_pct": 35.0,
    "fountain_min_operational_pct": 50.0,
    "paddle_min_operational_pct": 50.0,
    "manual_use_fountain": True,
    "manual_fountain_model": "B-501/B-503",
    "manual_fountain_qty": 1,
    "manual_use_paddlewheel": False,
    "manual_paddle_model": "B-209",
    "manual_paddle_qty": 0,
    "manual_use_radial": False,
    "manual_radial_model": "CRA-750 TS",
    "manual_radial_qty": 0,
    "manual_use_lobular": False,
    "manual_lobular_model": "Família SRT",
    "manual_lobular_qty": 0,
    "report_profile": "Produtor",
    "notes": "Projeto de tilápia em tanques revestidos superintensivos.",
}

for i, row in enumerate(DEFAULT_FEEDING_CURVE, start=1):
    sample_data[f"phase{i}_min_g"] = row["min_g"]
    sample_data[f"phase{i}_max_g"] = row["max_g"]
    sample_data[f"phase{i}_feeding_rate_percent"] = row["feeding_rate_percent"]
    sample_data[f"phase{i}_meals_per_day"] = row["meals_per_day"]
    sample_data[f"phase{i}_protein_percent"] = row["protein_percent"]
    sample_data[f"phase{i}_pellet_mm"] = row["pellet_mm"]
    sample_data[f"phase{i}_feed_price_per_kg"] = row["feed_price_per_kg"]
    sample_data[f"phase{i}_fcr"] = row.get("phase_fcr", 1.60)



def _is_streamlit_cloud_runtime() -> bool:
    """Detecta execução no Streamlit Community Cloud."""
    app_env = os.getenv("APP_ENV", "").strip().lower()
    cloud_flag = os.getenv("STREAMLIT_CLOUD", "").strip().lower()
    return (
        Path("/mount/src").exists()
        or app_env in {"cloud", "production", "prod"}
        or cloud_flag in {"1", "true", "yes"}
    )


def _supabase_required() -> bool:
    """Na nuvem, se a Supabase estiver configurada, o modo local fica bloqueado."""
    return is_supabase_configured() and _is_streamlit_cloud_runtime()


if "dash_form_data" not in st.session_state:
    st.session_state.dash_form_data = sample_data.copy()

if "current_project_id" not in st.session_state:
    st.session_state.current_project_id = None

if "selected_project_id" not in st.session_state:
    st.session_state.selected_project_id = None

if "latest_results" not in st.session_state:
    st.session_state.latest_results = None

if "storage_mode" not in st.session_state:
    st.session_state.storage_mode = "Supabase beta" if _supabase_required() else "Local"

if "auth_user_id" not in st.session_state:
    st.session_state.auth_user_id = None

if "auth_user_email" not in st.session_state:
    st.session_state.auth_user_email = None


def _merge_saved_inputs(saved_inputs: dict | None) -> dict:
    """Garante compatibilidade ao abrir projetos antigos com campos novos."""
    merged = sample_data.copy()
    if isinstance(saved_inputs, dict):
        merged.update(saved_inputs)
    return merged


def _use_supabase_storage() -> bool:
    return (
        st.session_state.get("storage_mode") == "Supabase beta"
        and bool(st.session_state.get("auth_user_id"))
    )


def _require_supabase_login() -> None:
    raise RuntimeError("Faça login na Supabase beta para usar projetos salvos na nuvem.")


def _list_projects_active() -> list[dict]:
    if st.session_state.get("storage_mode") == "Supabase beta":
        if _use_supabase_storage():
            return list_projects_remote(st.session_state["auth_user_id"])
        return []
    return list_projects()


def _load_project_active(project_id: str) -> dict:
    if st.session_state.get("storage_mode") == "Supabase beta":
        if _use_supabase_storage():
            return load_project_remote(project_id, st.session_state["auth_user_id"])
        _require_supabase_login()
    return load_project(project_id)


def _save_project_active(payload: dict, current_project_id: str | None = None) -> str:
    if st.session_state.get("storage_mode") == "Supabase beta":
        if _use_supabase_storage():
            return save_project_remote(
                payload,
                st.session_state["auth_user_id"],
                current_project_id=current_project_id,
            )
        _require_supabase_login()
    return save_project(payload)


def _delete_project_active(project_id: str) -> None:
    if st.session_state.get("storage_mode") == "Supabase beta":
        if _use_supabase_storage():
            delete_project_remote(project_id, st.session_state["auth_user_id"])
            return
        _require_supabase_login()
    delete_project(project_id)


def _duplicate_project_active(project_id: str) -> str:
    if st.session_state.get("storage_mode") == "Supabase beta":
        if _use_supabase_storage():
            return duplicate_project_remote(project_id, st.session_state["auth_user_id"])
        _require_supabase_login()
    return duplicate_project(project_id)


with st.sidebar:
    st.markdown("## Aqua Project Agent")
    st.markdown("Tilápia em tanques revestidos superintensivos")
    st.markdown("---")

    output_dir = st.text_input("Pasta de saída", "outputs")
    project_name_for_output = st.session_state.get("dash_form_data", {}).get("project_name", "Projeto tilápia tanques revestidos")
    suggested_output_name = _slugify_filename(project_name_for_output)
    output_name = st.text_input(
        "Nome base dos arquivos",
        value=suggested_output_name,
        help="Por padrão, o nome do arquivo acompanha o nome informado no campo Nome do projeto.",
    )

    if st.button("Carregar exemplo", key="btn_carregar_exemplo"):
        st.session_state.dash_form_data = sample_data.copy()
        st.session_state.current_project_id = None
        st.session_state.selected_project_id = None
        st.session_state.latest_results = None
        st.rerun()

    st.markdown("---")
    st.markdown("### Persistência")

    if _supabase_required():
        st.session_state.storage_mode = "Supabase beta"
        st.success("Modo de armazenamento: Supabase beta obrigatório na nuvem.")
        st.caption("Projetos salvos, abertura, duplicação e exclusão usam exclusivamente o banco Supabase.")
    elif is_supabase_configured():
        st.session_state.storage_mode = st.radio(
            "Modo de armazenamento",
            ["Local", "Supabase beta"],
            index=0 if st.session_state.get("storage_mode", "Local") == "Local" else 1,
            key="storage_mode_radio",
            help="Local é apenas para desenvolvimento no notebook. Na nuvem o sistema força Supabase beta.",
        )
    else:
        st.session_state.storage_mode = "Local"
        st.info("Supabase ainda não configurado no .env. Usando modo local.")

    if st.session_state.storage_mode == "Supabase beta":
        st.markdown("#### Login beta")
        email_login = st.text_input("Email do testador", key="sb_email_login")
        password_login = st.text_input("Senha", type="password", key="sb_password_login")
        c_auth1, c_auth2 = st.columns(2)

        with c_auth1:
            if st.button("Entrar", key="btn_supabase_entrar"):
                try:
                    auth_data = sign_in_with_password(email_login, password_login)
                    st.session_state.auth_user_id = auth_data.get("user_id")
                    st.session_state.auth_user_email = auth_data.get("email")
                    st.success(f"Login realizado: {st.session_state.auth_user_email}")
                    st.rerun()
                except Exception as exc:
                    st.error(f"Falha no login: {exc}")

        with c_auth2:
            if st.button("Sair", key="btn_supabase_sair"):
                st.session_state.auth_user_id = None
                st.session_state.auth_user_email = None
                st.session_state.current_project_id = None
                st.session_state.selected_project_id = None
                st.rerun()

        if st.session_state.get("auth_user_email"):
            st.markdown(
                f"""
<div class="connected-user-box">
    <div class="connected-user-label">Usuário conectado</div>
    <div class="connected-user-email">{st.session_state['auth_user_email']}</div>
</div>
""",
                unsafe_allow_html=True,
            )
        else:
            st.warning("Faça login para usar o armazenamento remoto.")

    st.markdown("---")
    st.markdown("### Projetos")

    try:
        projects = _list_projects_active()
    except Exception as exc:
        projects = []
        st.error(f"Não foi possível listar projetos: {exc}")

    project_options = [""] + [p["project_id"] for p in projects]
    project_labels = {
        "": "Selecione um projeto salvo",
        **{
            p["project_id"]: f"{p.get('project_name', 'Projeto sem nome')} ({str(p.get('updated_at', ''))[:19].replace('T', ' ')})"
            for p in projects
        },
    }

    if st.session_state.selected_project_id not in project_options:
        st.session_state.selected_project_id = None

    selected = st.selectbox(
        "Projetos salvos",
        project_options,
        format_func=lambda x: project_labels.get(x, x),
        index=project_options.index(st.session_state.selected_project_id)
        if st.session_state.selected_project_id in project_options
        else 0,
        key="saved_projects_select",
    )

    st.session_state.selected_project_id = selected if selected else None

    cproj1, cproj2 = st.columns(2)

    with cproj1:
        if st.button("Abrir projeto", key="btn_abrir_projeto"):
            if st.session_state.selected_project_id:
                try:
                    saved = _load_project_active(st.session_state.selected_project_id)
                    st.session_state.dash_form_data = _merge_saved_inputs(saved.get("inputs"))
                    st.session_state.current_project_id = saved.get("project_meta", {}).get("project_id")
                    st.session_state.latest_results = saved.get("results")
                    st.success("Projeto carregado com sucesso.")
                    st.rerun()
                except Exception as exc:
                    st.error(f"Erro ao abrir projeto: {exc}")

    with cproj2:
        if st.button("Novo projeto", key="btn_novo_projeto"):
            st.session_state.dash_form_data = sample_data.copy()
            st.session_state.current_project_id = None
            st.session_state.selected_project_id = None
            st.session_state.latest_results = None
            st.success("Novo projeto iniciado.")
            st.rerun()

    cproj3, cproj4 = st.columns(2)

    with cproj3:
        if st.button("Duplicar projeto", key="btn_duplicar_projeto"):
            if st.session_state.selected_project_id:
                try:
                    new_id = _duplicate_project_active(st.session_state.selected_project_id)
                    st.session_state.selected_project_id = new_id
                    st.session_state.current_project_id = new_id
                    st.success(f"Projeto duplicado: {new_id}")
                    st.rerun()
                except Exception as exc:
                    st.error(f"Erro ao duplicar projeto: {exc}")

    with cproj4:
        if st.button("Excluir projeto", key="btn_excluir_projeto"):
            if st.session_state.selected_project_id:
                try:
                    _delete_project_active(st.session_state.selected_project_id)
                    st.session_state.current_project_id = None
                    st.session_state.selected_project_id = None
                    st.session_state.latest_results = None
                    st.success("Projeto excluído.")
                    st.rerun()
                except Exception as exc:
                    st.error(f"Erro ao excluir projeto: {exc}")

    if st.session_state.current_project_id:
        st.caption(f"Projeto atual: {st.session_state.current_project_id}")

    st.markdown("---")
    st.markdown("### Uso")
    st.markdown(
        """
- Ajuste os dados técnicos
- Revise a alimentação por fase
- Confira o dashboard
- Gere o projeto completo
- Salve o projeto com resultados
- Na nuvem, use login Supabase beta para salvar e abrir projetos
"""
    )

fd = st.session_state.dash_form_data
report_profile = fd.get("report_profile", "Produtor")

st.markdown(
    """
<div class="hero-box">
    <div class="hero-title">Aqua Project Agent Dashboard</div>
    <div class="hero-subtitle">Sistema especialista em tilápia para tanques revestidos superintensivos</div>
</div>
""",
    unsafe_allow_html=True,
)

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(
    [
        "Projeto básico",
        "Crescimento",
        "Alimentação avançada",
        "Oxigênio e aeração",
        "Dashboard",
        "Relatório Profissional",
    ]
)


with tab1:
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    st.subheader("Projeto básico")

    fd["report_profile"] = st.selectbox(
        "Perfil do relatório",
        ["Produtor", "Técnico", "Banco/Financiamento"],
        index=["Produtor", "Técnico", "Banco/Financiamento"].index(fd.get("report_profile", "Produtor")),
        key="report_profile_select",
    )
    report_profile = fd["report_profile"]

    fd["system_type"] = st.selectbox(
        "Tipo de estrutura",
        ["Circular revestido", "Suspenso revestido retangular", "Escavado revestido"],
        index=["Circular revestido", "Suspenso revestido retangular", "Escavado revestido"].index(
            fd.get("system_type", "Circular revestido") if fd.get("system_type", "Circular revestido") in ["Circular revestido", "Suspenso revestido retangular", "Escavado revestido"] else "Circular revestido"
        ),
    )
    suggested_depth = recommended_depth_for_structure(fd["system_type"])

    g1, g2, g3 = st.columns(3)
    with g1:
        fd["project_name"] = st.text_input("Nome do projeto", fd["project_name"])
        fd["author_name"] = st.text_input("Autor", fd["author_name"])
        st.text_input("Espécie", value="Tilápia", disabled=True)
        fd["species"] = "Tilápia"
        fd["region_focus"] = st.text_input("Região de referência", fd["region_focus"])
        fd["number_of_units"] = st.number_input("Tanques informados", min_value=1, value=int(fd["number_of_units"]), step=1)

    with g2:
        fd["initial_weight_g"] = st.number_input("Peso inicial (g)", min_value=1.0, value=float(max(fd.get("initial_weight_g", 1.0), 1.0)), step=1.0)
        fd["target_weight_g"] = st.number_input("Peso final alvo (g)", min_value=100.0, max_value=2000.0, value=float(min(max(fd.get("target_weight_g", 1000.0), 100.0), 2000.0)), step=10.0)
        fd["density_kg_m3"] = st.number_input("Densidade final (kg/m³)", min_value=10.0, max_value=50.0, value=float(min(max(fd.get("density_kg_m3", 30.0), 10.0), 50.0)), step=0.5)
        fd["survival_rate"] = st.number_input("Sobrevivência (0 a 1)", min_value=0.01, max_value=0.999, value=float(fd.get("survival_rate", 0.90)), step=0.01)
        fd["water_temperature_c"] = st.number_input("Temperatura média da água (°C)", min_value=10.0, max_value=40.0, value=float(fd.get("water_temperature_c", 28.0)), step=0.5)

    with g3:
        fd["production_strategy"] = st.selectbox("Estratégia de produção", ["Ciclos simultâneos", "Escalonada"], index=0 if fd.get("production_strategy", "Ciclos simultâneos") == "Ciclos simultâneos" else 1)
        if fd["production_strategy"] == "Escalonada":
            fd["scheduling_basis"] = st.selectbox("Critério de escalonamento", ["Intervalo entre despescas", "Número de lotes", "Tanques por lote"], index=["Intervalo entre despescas", "Número de lotes", "Tanques por lote"].index(fd.get("scheduling_basis", "Intervalo entre despescas")))
            if fd["scheduling_basis"] == "Intervalo entre despescas":
                fd["harvest_interval_months"] = st.number_input("Intervalo desejado entre despescas (meses)", min_value=0.25, value=float(fd.get("harvest_interval_months", 1.0)), step=0.25)
            elif fd["scheduling_basis"] == "Número de lotes":
                fd["manual_parallel_batches"] = st.number_input("Número de lotes desejado", min_value=1, value=int(fd.get("manual_parallel_batches", max(fd.get("number_of_units", 1), 1))), step=1)
            else:
                fd["desired_units_per_batch"] = st.number_input("Tanques por lote", min_value=1, value=int(fd.get("desired_units_per_batch", 1)), step=1)
        else:
            fd["scheduling_basis"] = "Intervalo entre despescas"
        fd["sale_price_per_kg"] = st.number_input("Preço de venda (R$/kg)", min_value=0.01, value=float(fd.get("sale_price_per_kg", 10.0)), step=0.1)
        fd["fingerling_price"] = st.number_input("Preço do alevino/recria (R$/unidade)", min_value=0.0, value=float(fd.get("fingerling_price", 0.30)), step=0.05)

    st.markdown("#### Geometria da unidade")
    h1, h2, h3 = st.columns(3)
    if fd["system_type"] == "Circular revestido":
        with h1:
            fd["unit_volume_m3"] = st.number_input("Volume útil por unidade (m³)", min_value=1.0, value=float(fd["unit_volume_m3"]), step=1.0)
        with h2:
            fd["water_depth_m"] = st.number_input("Altura da água (m)", min_value=0.5, value=float(fd.get("water_depth_m", suggested_depth)), step=0.1, help=f"Sugestão do sistema para este tipo de estrutura: {suggested_depth:.1f} m")
        with h3:
            st.caption("Para tanques circulares, o diâmetro é calculado automaticamente a partir do volume e da profundidade.")
    else:
        st.info("Para estruturas não circulares, o sistema considera comprimento, largura e profundidade para validar o volume útil e a admissibilidade do chafariz.")
        with h1:
            fd["tank_length_m"] = st.number_input("Comprimento interno (m)", min_value=1.0, value=float(fd.get("tank_length_m", 10.0)), step=0.5)
        with h2:
            fd["tank_width_m"] = st.number_input("Largura interna (m)", min_value=1.0, value=float(fd.get("tank_width_m", 10.0)), step=0.5)
        with h3:
            fd["water_depth_m"] = st.number_input("Altura da água (m)", min_value=0.5, value=float(fd.get("water_depth_m", suggested_depth)), step=0.1, help=f"Sugestão do sistema para este tipo de estrutura: {suggested_depth:.1f} m")

    geom_ui = _ui_geometry(fd["system_type"], float(fd.get("unit_volume_m3", 0.0)), float(fd.get("tank_length_m", 0.0)), float(fd.get("tank_width_m", 0.0)), float(fd.get("water_depth_m", suggested_depth)))
    fd["unit_volume_m3"] = geom_ui["volume_m3"]
    m1, m2, m3 = st.columns(3)
    with m1:
        if fd["system_type"] == "Circular revestido":
            st.metric("Diâmetro interno estimado", f"{num(geom_ui['diameter_m'])} m")
        else:
            st.metric("Comprimento x largura", f"{num(geom_ui['length_m'])} x {num(geom_ui['width_m'])} m")
    with m2:
        st.metric("Área superficial por tanque", f"{num(geom_ui['area_m2'])} m²")
    with m3:
        st.metric("Volume útil calculado", f"{num(geom_ui['volume_m3'])} m³")

    st.markdown("#### Modelo econômico")
    e1, e2 = st.columns(2)
    with e1:
        fd["economic_model_mode"] = st.selectbox("Modelo econômico", ["Simplificado", "Fixo + por tanque", "Composição preliminar", "Manual (CAPEX total)"], index=["Simplificado", "Fixo + por tanque", "Composição preliminar", "Manual (CAPEX total)"].index(fd.get("economic_model_mode", "Simplificado") if fd.get("economic_model_mode", "Simplificado") in ["Simplificado", "Fixo + por tanque", "Composição preliminar", "Manual (CAPEX total)"] else "Simplificado"))
    with e2:
        if fd["economic_model_mode"] == "Simplificado":
            fd["cost_scaling_mode"] = st.selectbox("Escala dos custos simplificados", ["Fixos (não escalar)", "Escalonar pelo nº de tanques"], index=["Fixos (não escalar)", "Escalonar pelo nº de tanques"].index(fd.get("cost_scaling_mode", "Fixos (não escalar)") if fd.get("cost_scaling_mode", "Fixos (não escalar)") in ["Fixos (não escalar)", "Escalonar pelo nº de tanques"] else "Fixos (não escalar)"))
            fd["cost_reference_units"] = st.number_input("Número de unidades de referência", min_value=1, value=int(max(fd.get("cost_reference_units", fd["number_of_units"]), 1)), step=1)

    ec1, ec2 = st.columns(2)
    with ec1:
        if fd["economic_model_mode"] == "Simplificado":
            fd["electricity_cost_cycle"] = st.number_input("Outros custos de energia por ciclo (R$)", min_value=0.0, value=float(fd.get("electricity_cost_cycle", 0.0)), step=500.0)
            fd["labor_cost_cycle"] = st.number_input("Custo de mão de obra por ciclo (R$)", min_value=0.0, value=float(fd.get("labor_cost_cycle", 0.0)), step=500.0)
            fd["other_costs_cycle"] = st.number_input("Outros custos por ciclo (R$)", min_value=0.0, value=float(fd.get("other_costs_cycle", 0.0)), step=500.0)
        elif fd["economic_model_mode"] == "Fixo + por tanque":
            fd["electricity_cost_fixed_cycle"] = st.number_input("Energia fixa por ciclo (R$)", min_value=0.0, value=float(fd.get("electricity_cost_fixed_cycle", 0.0)), step=500.0)
            fd["electricity_cost_per_unit_cycle"] = st.number_input("Energia por tanque por ciclo (R$)", min_value=0.0, value=float(fd.get("electricity_cost_per_unit_cycle", 0.0)), step=100.0)
            fd["labor_cost_fixed_cycle"] = st.number_input("Mão de obra fixa por ciclo (R$)", min_value=0.0, value=float(fd.get("labor_cost_fixed_cycle", 0.0)), step=500.0)
            fd["labor_cost_per_unit_cycle"] = st.number_input("Mão de obra por tanque por ciclo (R$)", min_value=0.0, value=float(fd.get("labor_cost_per_unit_cycle", 0.0)), step=100.0)
        else:
            fd["electricity_cost_cycle"] = st.number_input("Outros custos de energia por ciclo (R$)", min_value=0.0, value=float(fd.get("electricity_cost_cycle", 0.0)), step=500.0)
            fd["labor_cost_cycle"] = st.number_input("Custo de mão de obra por ciclo (R$)", min_value=0.0, value=float(fd.get("labor_cost_cycle", 0.0)), step=500.0)
            fd["other_costs_cycle"] = st.number_input("Outros custos por ciclo (R$)", min_value=0.0, value=float(fd.get("other_costs_cycle", 0.0)), step=500.0)
    with ec2:
        if fd["economic_model_mode"] == "Simplificado":
            fd["capex_total"] = st.number_input("CAPEX total (R$)", min_value=0.0, value=float(fd.get("capex_total", 0.0)), step=1000.0)
        elif fd["economic_model_mode"] == "Fixo + por tanque":
            fd["other_cost_fixed_cycle"] = st.number_input("Outros custos fixos por ciclo (R$)", min_value=0.0, value=float(fd.get("other_cost_fixed_cycle", 0.0)), step=500.0)
            fd["other_cost_per_unit_cycle"] = st.number_input("Outros custos por tanque por ciclo (R$)", min_value=0.0, value=float(fd.get("other_cost_per_unit_cycle", 0.0)), step=100.0)
            fd["capex_fixed_total"] = st.number_input("CAPEX fixo total (R$)", min_value=0.0, value=float(fd.get("capex_fixed_total", 0.0)), step=1000.0)
            fd["capex_per_unit"] = st.number_input("CAPEX por tanque (R$)", min_value=0.0, value=float(fd.get("capex_per_unit", 0.0)), step=500.0)
        elif fd["economic_model_mode"] == "Composição preliminar":
            capex_preview = _estimate_preliminary_capex_from_form(fd, geom_ui)
            fd["capex_total"] = float(capex_preview["CAPEX preliminar total"])
            st.metric("CAPEX preliminar estimado", brl(fd["capex_total"]))
            st.caption("Estimativa preliminar por composição. Atualize com cotações reais antes de usar para financiamento ou decisão de investimento.")
            capex_df = pd.DataFrame([
                {"Componente": k, "Valor estimado": brl(v)}
                for k, v in capex_preview.items()
                if isinstance(v, (int, float)) and k != "CAPEX preliminar total"
            ])
            st.dataframe(capex_df, use_container_width=True, hide_index=True)
        else:
            fd["capex_total"] = st.number_input("CAPEX total manual (R$)", min_value=0.0, value=float(fd.get("capex_total", 0.0)), step=1000.0)
        fd["notes"] = st.text_area("Observações", fd["notes"], height=120)

    if fd.get("economic_model_mode") == "Simplificado":
        ref_units = max(1.0, float(fd.get("cost_reference_units", fd.get("number_of_units", 1)) or 1))
        current_units = max(1.0, float(fd.get("number_of_units", 1) or 1))
        scale_factor = current_units / ref_units if fd.get("cost_scaling_mode") == "Escalonar pelo nº de tanques" else 1.0
        st.markdown("#### Prévia dos custos simplificados aplicados")
        st.caption(
            "Esta prévia mostra o efeito da opção de escala dos custos. "
            "Em 'Fixos', os valores digitados são usados como estão. "
            "Em 'Escalonar pelo nº de tanques', os custos são multiplicados pelo fator tanques informados / unidades de referência."
        )
        pc1, pc2, pc3, pc4 = st.columns(4)
        with pc1:
            st.metric("Fator aplicado", f"{num(scale_factor, 2)}x")
        with pc2:
            st.metric("Energia/ciclo aplicada", brl(float(fd.get("electricity_cost_cycle", 0.0)) * scale_factor))
        with pc3:
            st.metric("Mão de obra/ciclo aplicada", brl(float(fd.get("labor_cost_cycle", 0.0)) * scale_factor))
        with pc4:
            st.metric("Outros custos/ciclo aplicados", brl(float(fd.get("other_costs_cycle", 0.0)) * scale_factor))

    st.caption("A altura da água é sugerida automaticamente conforme o tipo de estrutura, mas pode ser alterada pelo usuário. Alterações na geometria mudam volume útil, biomassa por tanque, produção e aeração.")
    st.markdown("</div>", unsafe_allow_html=True)

with tab2:
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    st.subheader("Crescimento e temperatura")

    c1, c2 = st.columns(2)

    with c1:
        st.metric("Peso inicial do projeto", f"{num(fd['initial_weight_g'])} g")
        st.metric("Peso final alvo do projeto", f"{num(fd['target_weight_g'])} g")
        st.metric("Temperatura média da água", f"{num(fd['water_temperature_c'])} °C")
        st.caption("Esses valores são definidos na aba Projeto básico e alimentam automaticamente o cálculo biológico do ciclo.")

    with c2:
        fd["growth_curve_adjustment_pct"] = st.slider(
            "Ajuste global da curva de crescimento (%)",
            min_value=80,
            max_value=120,
            value=int(fd.get("growth_curve_adjustment_pct", 100.0)),
            step=1,
        )
        st.info("A curva de crescimento consolidada usa subfaixas de peso a 28 °C: 1–5 g, 5–20 g, 20–100 g, 100–250 g, 250–650 g, 650–1.000 g e 1.000–2.000 g. A tabela alimentar trabalha com 5 fases práticas: 1–20 g, 20–250 g, 250–650 g, 650–1.000 g e 1.000–2.000 g.")
        st.caption("Referências consolidadas: Embrapa, Auburn/Rakocy e literatura técnica para tilápia intensiva em tanques.")

    st.markdown("</div>", unsafe_allow_html=True)

with tab3:
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    st.subheader("Plano alimentar por fases")
    st.caption("Edite o FCR por fase aqui, junto com os demais parâmetros da alimentação. A duração de cada fase é calculada automaticamente a partir do peso inicial, peso final, temperatura e ajuste global da curva de crescimento. A tabela-resumo mostrará a granulometria recomendada por faixa dentro de cada fase.")

    for i in range(1, 6):
        with st.expander(f"Fase {i}", expanded=(i == 1)):
            c1, c2, c3, c4 = st.columns(4)

            with c1:
                fd[f"phase{i}_min_g"] = st.number_input(
                    f"Peso mín. fase {i} (g)",
                    min_value=0.0,
                    value=float(fd[f"phase{i}_min_g"]),
                    step=1.0,
                    key=f"min_{i}",
                )
                fd[f"phase{i}_feeding_rate_percent"] = st.number_input(
                    f"Arraçoamento fase {i} (%)",
                    min_value=0.1,
                    value=float(fd[f"phase{i}_feeding_rate_percent"]),
                    step=0.1,
                    key=f"feedr_{i}",
                )

            with c2:
                fd[f"phase{i}_max_g"] = st.number_input(
                    f"Peso máx. fase {i} (g)",
                    min_value=1.0,
                    value=float(fd[f"phase{i}_max_g"]),
                    step=1.0,
                    key=f"max_{i}",
                )
                fd[f"phase{i}_meals_per_day"] = st.number_input(
                    f"Tratos/dia fase {i}",
                    min_value=1,
                    value=int(fd[f"phase{i}_meals_per_day"]),
                    step=1,
                    key=f"meals_{i}",
                )

            with c3:
                fd[f"phase{i}_protein_percent"] = st.number_input(
                    f"Proteína fase {i} (%)",
                    min_value=10.0,
                    max_value=60.0,
                    value=float(fd[f"phase{i}_protein_percent"]),
                    step=0.5,
                    key=f"prot_{i}",
                )
                fd[f"phase{i}_pellet_mm"] = st.number_input(
                    f"Granulometria base fase {i} (mm)",
                    min_value=0.1,
                    value=float(fd[f"phase{i}_pellet_mm"]),
                    step=0.1,
                    key=f"pel_{i}",
                )

            with c4:
                fd[f"phase{i}_feed_price_per_kg"] = st.number_input(
                    f"Preço ração fase {i} (R$/kg)",
                    min_value=0.01,
                    value=float(fd[f"phase{i}_feed_price_per_kg"]),
                    step=0.1,
                    key=f"price_{i}",
                )

                fd[f"phase{i}_fcr"] = st.number_input(
                    f"FCR fase {i}",
                    min_value=0.5,
                    max_value=3.0,
                    value=float(fd.get(f"phase{i}_fcr", [1.10, 1.35, 1.55, 1.65, 1.75][i - 1])),
                    step=0.01,
                    key=f"fcr_{i}",
                )

    st.markdown("</div>", unsafe_allow_html=True)


with tab4:
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    st.subheader("Oxigênio e aeração")

    st.info("A aeração é obrigatória para tilápia em tanques revestidos superintensivos. O dimensionamento considera demanda de O₂ do cultivo, temperatura, altitude, eficiência de campo e fator de segurança. O sistema também estima a operação por fase para reduzir custo energético sem perder segurança operacional.")

    q1, q2, q3, q4 = st.columns(4)
    with q1:
        st.metric("OD alvo", "5,0–7,5 mg/L")
    with q2:
        st.metric("Alerta de desempenho", "< 3,5 mg/L")
    with q3:
        st.metric("pH operacional", "6,5–8,5")
    with q4:
        st.metric("Alcalinidade", "100–250 mg/L")

    st.caption("Faixas de referência consolidadas para qualidade de água: CO₂ < 40 mg/L; NH₃-N com atenção a partir de 1,0 mg/L e risco severo acima de 2,0 mg/L.")

    c1, c2 = st.columns(2)
    with c1:
        fd["oxygen_demand_mg_per_kg_h"] = st.number_input(
            "Demanda-base de oxigênio (mg O₂/kg/h)",
            min_value=100.0,
            value=float(fd.get("oxygen_demand_mg_per_kg_h", 550.0)),
            step=10.0,
        )
        fd["site_altitude_m"] = st.number_input(
            "Altitude do local (m)",
            min_value=0.0,
            value=float(fd.get("site_altitude_m", 0.0)),
            step=50.0,
        )
        fd["field_efficiency_pct"] = st.number_input(
            "Eficiência de campo (%)",
            min_value=10.0,
            max_value=100.0,
            value=float(fd.get("field_efficiency_pct", 85.0)),
            step=1.0,
        )
        fd["aeration_safety_factor_pct"] = st.number_input(
            "Fator de segurança da aeração (%)",
            min_value=0.0,
            max_value=100.0,
            value=float(fd.get("aeration_safety_factor_pct", 20.0)),
            step=1.0,
        )
    with c2:
        fd["target_do_mg_l"] = st.number_input(
            "OD alvo operacional (mg/L)",
            min_value=3.0,
            max_value=10.0,
            value=float(fd.get("target_do_mg_l", 5.0)),
            step=0.1,
        )
        fd["aeration_hours_per_day"] = st.number_input(
            "Horas de aeração por dia",
            min_value=0.0,
            max_value=24.0,
            value=float(fd.get("aeration_hours_per_day", 24.0)),
            step=0.5,
        )
        fd["electricity_price_kwh"] = st.number_input(
            "Preço da energia (R$/kWh)",
            min_value=0.01,
            value=float(fd.get("electricity_price_kwh", 0.45)),
            step=0.05,
        )

    current_altitude_factor = altitude_transfer_factor(fd.get("site_altitude_m", 0.0))
    if float(fd.get("site_altitude_m", 0.0) or 0.0) >= 4000:
        st.error(
            "Altitude extremamente elevada para o cenário aquícola. O sistema penaliza fortemente a transferência "
            "de oxigênio e recomenda tratar a simulação como operacionalmente crítica."
        )
    elif float(fd.get("site_altitude_m", 0.0) or 0.0) >= 2500:
        st.warning(
            "Altitude elevada: a eficiência de transferência de oxigênio foi reduzida no cálculo. "
            "Valide equipamentos e manejo com dados locais."
        )
    st.caption(
        f"Fator operacional de correção por altitude aplicado na aeração: {num(current_altitude_factor, 2)}x."
    )

    fd["aeration_power_mode"] = st.selectbox(
        "Modo de operação da aeração ao longo do ciclo",
        ["Potência modulada por fase", "Potência fixa no ciclo inteiro"],
        index=0 if fd.get("aeration_power_mode", "Potência modulada por fase") == "Potência modulada por fase" else 1,
    )
    if fd["aeration_power_mode"] == "Potência modulada por fase":
        fd["aeration_control_strategy"] = st.selectbox(
            "Estratégia de controle da potência",
            ["Automático", "Inversor de frequência", "Acionamento por etapas", "Híbrido"],
            index=["Automático", "Inversor de frequência", "Acionamento por etapas", "Híbrido"].index(fd.get("aeration_control_strategy", "Automático") if fd.get("aeration_control_strategy", "Automático") in ["Automático", "Inversor de frequência", "Acionamento por etapas", "Híbrido"] else "Automático"),
        )
        p1, p2, p3 = st.columns(3)
        with p1:
            fd["blower_min_operational_pct"] = st.number_input("Mínimo operacional sopradores (%)", min_value=10.0, max_value=100.0, value=float(fd.get("blower_min_operational_pct", 35.0)), step=1.0)
        with p2:
            fd["fountain_min_operational_pct"] = st.number_input("Mínimo operacional chafariz (%)", min_value=10.0, max_value=100.0, value=float(fd.get("fountain_min_operational_pct", 50.0)), step=1.0)
        with p3:
            fd["paddle_min_operational_pct"] = st.number_input("Mínimo operacional pás (%)", min_value=10.0, max_value=100.0, value=float(fd.get("paddle_min_operational_pct", 50.0)), step=1.0)

    fd["aeration_mode"] = st.selectbox(
        "Modo de aeração",
        ["Automático", "Manual"],
        index=0 if fd.get("aeration_mode", "Automático") == "Automático" else 1,
    )

    if fd["aeration_mode"] == "Automático":
        aeration_tech_options = [
            "Chafariz",
            "Pás",
            "Soprador",
            "Híbrido: Chafariz + Pás",
            "Híbrido: Soprador + Pás",
            "Híbrido: Soprador + Chafariz",
            "Híbrido: Soprador + Chafariz + Pás",
        ]
        current_aeration_tech = fd.get("automatic_aeration_technology", "Chafariz")
        if current_aeration_tech not in aeration_tech_options:
            current_aeration_tech = "Chafariz"
        fd["automatic_aeration_technology"] = st.selectbox(
            "Tecnologia automática",
            aeration_tech_options,
            index=aeration_tech_options.index(current_aeration_tech),
            help=(
                "Use as opções híbridas quando fizer sentido combinar oxigenação superficial, difusão por soprador "
                "e circulação/vórtice para concentração de sólidos."
            ),
        )
        if fd["automatic_aeration_technology"] in ("Chafariz", "Pás", "Híbrido: Chafariz + Pás"):
            st.caption("A admissibilidade agora considera o diâmetro/área do tanque. Tanques circulares grandes podem aceitar arranjos A1–A4 de chafarizes e também aeradores de pás como apoio à circulação/vórtice.")
        if "Soprador" in fd["automatic_aeration_technology"]:
            st.caption("Para sopradores, a simulação estima também a metragem preliminar de mangueira porosa/difusor por tanque. Ajuste com dados reais do fabricante.")
            fd["blower_type"] = st.selectbox(
                "Tipo de soprador",
                ["Automático", "Radial", "Lobular"],
                index=["Automático", "Radial", "Lobular"].index(fd.get("blower_type", "Automático")),
            )
            fd["diffusion_efficiency_pct"] = st.number_input(
                "Eficiência efetiva de transferência do sistema de difusão (%) — base para OTR/m",
                min_value=1.0,
                max_value=100.0,
                value=float(fd.get("diffusion_efficiency_pct", 12.0)),
                step=1.0,
            )
            st.caption(
                "12% permanece como premissa conservadora/editável para mangueira porosa BERAQUA sem SOTR/SAE certificado. "
                "O cálculo da metragem usa OTR preliminar por metro: conservador, padrão, otimista ou alta eficiência conforme este campo."
            )
    else:
        st.caption("No modo manual, o sistema permite combinar tecnologias. A recomendação automática continuará visível abaixo como referência técnica.")
        m1, m2 = st.columns(2)
        with m1:
            fd["manual_use_fountain"] = st.checkbox("Ativar chafariz", value=bool(fd.get("manual_use_fountain", True)))
            if fd["manual_use_fountain"]:
                fountain_models = _model_options(FOUNTAIN_LIBRARY)
                current = fd.get("manual_fountain_model", fountain_models[0])
                fd["manual_fountain_model"] = st.selectbox("Modelo de chafariz", fountain_models, index=fountain_models.index(current) if current in fountain_models else 0)
                fd["manual_fountain_qty"] = st.number_input("Quantidade de chafarizes", min_value=0, value=int(fd.get("manual_fountain_qty", 1)), step=1)

            fd["manual_use_paddlewheel"] = st.checkbox("Ativar aerador de pás", value=bool(fd.get("manual_use_paddlewheel", False)))
            if fd["manual_use_paddlewheel"]:
                paddle_models = _model_options(PADDLEWHEEL_LIBRARY)
                current = fd.get("manual_paddle_model", paddle_models[0])
                fd["manual_paddle_model"] = st.selectbox("Modelo de pás", paddle_models, index=paddle_models.index(current) if current in paddle_models else 0)
                fd["manual_paddle_qty"] = st.number_input("Quantidade de aeradores de pás", min_value=0, value=int(fd.get("manual_paddle_qty", 1)), step=1)

        with m2:
            fd["manual_use_radial"] = st.checkbox("Ativar soprador radial", value=bool(fd.get("manual_use_radial", False)))
            if fd["manual_use_radial"]:
                radial_models = _model_options(RADIAL_BLOWER_LIBRARY)
                current = fd.get("manual_radial_model", radial_models[0])
                fd["manual_radial_model"] = st.selectbox("Modelo de soprador radial", radial_models, index=radial_models.index(current) if current in radial_models else 0)
                fd["manual_radial_qty"] = st.number_input("Quantidade de sopradores radiais", min_value=0, value=int(fd.get("manual_radial_qty", 1)), step=1)
                fd["diffusion_efficiency_pct"] = st.number_input("Eficiência efetiva de transferência do sistema de difusão (%) — base para OTR/m", min_value=1.0, max_value=100.0, value=float(fd.get("diffusion_efficiency_pct", 12.0)), step=1.0, key="diff_radial")

            fd["manual_use_lobular"] = st.checkbox("Ativar soprador lobular", value=bool(fd.get("manual_use_lobular", False)))
            if fd["manual_use_lobular"]:
                lobular_models = _model_options(LOBULAR_BLOWER_LIBRARY)
                current = fd.get("manual_lobular_model", lobular_models[0])
                fd["manual_lobular_model"] = st.selectbox("Modelo de soprador lobular", lobular_models, index=lobular_models.index(current) if current in lobular_models else 0)
                fd["manual_lobular_qty"] = st.number_input("Quantidade de sopradores lobulares", min_value=0, value=int(fd.get("manual_lobular_qty", 1)), step=1)
                fd["diffusion_efficiency_pct"] = st.number_input("Eficiência efetiva de transferência do sistema de difusão (%) — base para OTR/m", min_value=1.0, max_value=100.0, value=float(fd.get("diffusion_efficiency_pct", 12.0)), step=1.0, key="diff_lobular")

    st.markdown("</div>", unsafe_allow_html=True)


def _normalize_production_schedule_for_display(results: dict, form_data: dict) -> dict:
    """
    Corrige/normaliza, no nível do app, a distribuição operacional de lotes
    para evitar inconsistência visual quando o cálculo informa lotes em paralelo,
    mas o campo "Tanques por lote" fica desatualizado.

    Exemplo esperado:
    18 tanques / 9 lotes em paralelo = 2 tanques por lote.
    """
    if not isinstance(results, dict):
        return {}

    base = results.get("base", {})
    if not isinstance(base, dict):
        return {}

    schedule = base.get("production_schedule", {})
    if not isinstance(schedule, dict):
        return {}

    try:
        number_of_units = max(1, int(form_data.get("number_of_units", 1)))
    except (TypeError, ValueError):
        number_of_units = 1

    strategy = schedule.get("production_strategy", form_data.get("production_strategy", "Ciclos simultâneos"))
    if strategy != "Escalonada":
        schedule["batches_in_parallel"] = 1
        schedule["units_per_batch_exact"] = float(number_of_units)
        schedule["batch_distribution_note"] = (
            f"No modo de ciclos simultâneos, os {number_of_units} tanque(s) "
            "entram e saem do ciclo no mesmo período."
        )
        return schedule

    scheduling_basis = form_data.get("scheduling_basis", schedule.get("scheduling_basis", "Intervalo entre despescas"))

    estimated_cycle_days = base.get("estimated_cycle_days", 0) or 0
    try:
        cycle_months = float(schedule.get("first_harvest_after_months") or (float(estimated_cycle_days) / 30.4375))
    except (TypeError, ValueError, ZeroDivisionError):
        cycle_months = 0.0

    try:
        harvest_interval = float(form_data.get("harvest_interval_months", schedule.get("harvest_interval_months", 1.0)) or 1.0)
    except (TypeError, ValueError):
        harvest_interval = 1.0
    harvest_interval = max(harvest_interval, 0.01)

    def _ceil_positive(value: float) -> int:
        if value <= 1:
            return 1
        as_int = int(value)
        return as_int if abs(value - as_int) < 1e-9 else as_int + 1

    if scheduling_basis == "Número de lotes":
        try:
            batches = max(1, int(form_data.get("manual_parallel_batches", schedule.get("batches_in_parallel", 1)) or 1))
        except (TypeError, ValueError):
            batches = 1
        batch_mode = "Personalizado"
    elif scheduling_basis == "Tanques por lote":
        try:
            desired_units = max(1, int(form_data.get("desired_units_per_batch", 1) or 1))
        except (TypeError, ValueError):
            desired_units = 1
        batches = _ceil_positive(number_of_units / desired_units)
        batch_mode = "Personalizado"
    else:
        batches = _ceil_positive(cycle_months / harvest_interval) if cycle_months > 0 else max(1, int(schedule.get("batches_in_parallel", 1) or 1))
        batch_mode = "Automático"

    batches = max(1, min(int(batches), number_of_units))
    units_per_batch_exact = number_of_units / batches

    schedule["production_strategy"] = "Escalonada"
    schedule["scheduling_basis"] = scheduling_basis
    schedule["harvest_interval_months"] = harvest_interval
    schedule["batches_in_parallel"] = batches
    schedule["units_per_batch_exact"] = units_per_batch_exact
    schedule["production_batch_mode"] = schedule.get("production_batch_mode", batch_mode)

    if abs(units_per_batch_exact - round(units_per_batch_exact)) < 1e-9:
        units_text = f"{int(round(units_per_batch_exact))} tanque(s) por lote"
    else:
        units_text = f"{units_per_batch_exact:.2f} tanque(s) por lote em média"

    schedule["batch_distribution_note"] = (
        f"Com {number_of_units} tanque(s) e {batches} lote(s) em paralelo, "
        f"a distribuição operacional sugerida é de {units_text}; "
        "a primeira receita ocorre apenas após o fechamento do ciclo biológico inicial."
    )

    return schedule

fd_calc = fd.copy()
fd_calc["fingerling_weight_kg"] = fd["fingerling_weight_kg"] / 1000.0
# O modelo de dados ainda não tem um modo econômico específico para composição;
# por isso, a composição preliminar calcula capex_total na interface e envia ao motor
# como CAPEX manual, preservando o valor estimado.
if fd_calc.get("economic_model_mode") == "Composição preliminar":
    fd_calc["economic_model_mode"] = "Manual (CAPEX total)"

allowed_input_fields = {field.name for field in fields(DashboardProjectInput)}
fd_model = {key: value for key, value in fd_calc.items() if key in allowed_input_fields}
ignored_input_fields = sorted(set(fd_calc) - allowed_input_fields)
if ignored_input_fields:
    st.session_state["ignored_input_fields_for_model"] = ignored_input_fields

inp = DashboardProjectInput(**fd_model)
results = calculate_dashboard_project(inp)
st.session_state.latest_results = results
base_res = results["base"]
_normalize_production_schedule_for_display(results, fd)
base_res = results["base"]
altitude_correction_ui = {
    "altitude_m": fd_calc.get("site_altitude_m", 0.0),
    "altitude_factor": base_res.get("altitude_transfer_factor"),
    "installed_supply_original": base_res.get("installed_oxygen_supply_sea_level_kg_h"),
    "installed_supply_corrected": base_res.get("installed_oxygen_supply_kg_h"),
    "oxygen_offer_demand_ratio": base_res.get("oxygen_offer_demand_ratio"),
}
base_res = results["base"]
alt_factor_ui = altitude_transfer_factor(fd_calc.get("site_altitude_m", 0.0))
temp_factor_ui = surface_aeration_base_factor(fd_calc.get("water_temperature_c", 28.0))


with tab2:
    feeding_plan_calc = pd.DataFrame(base_res.get("feeding_plan", []))
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    st.subheader("Resultado da curva de crescimento")
    gcalc1, gcalc2, gcalc3 = st.columns(3)
    with gcalc1:
        st.metric("Ciclo biológico estimado", f"{num(base_res.get('estimated_cycle_days'), 0)} dias")
    with gcalc2:
        total_gain_g = max(float(fd.get("target_weight_g", 0.0)) - float(fd.get("initial_weight_g", 0.0)), 0.0)
        cycle_days = float(base_res.get("estimated_cycle_days", 0.0) or 0.0)
        avg_daily_gain = total_gain_g / cycle_days if cycle_days > 0 else 0.0
        st.metric("Ganho médio estimado", f"{num(avg_daily_gain, 2)} g/peixe/dia")
    with gcalc3:
        st.metric("Ajuste aplicado", f"{num(fd.get('growth_curve_adjustment_pct', 100.0), 0)}%")

    if not feeding_plan_calc.empty:
        growth_cols = [
            "phase_name",
            "weight_range",
            "days_in_phase",
            "phase_daily_growth_g",
            "phase_fcr",
        ]
        growth_view = feeding_plan_calc[[c for c in growth_cols if c in feeding_plan_calc.columns]].copy()
        growth_view = growth_view.rename(
            columns={
                "phase_name": "Fase",
                "weight_range": "Faixa de peso",
                "days_in_phase": "Duração estimada (dias)",
                "phase_daily_growth_g": "Ganho médio da fase (g/dia)",
                "phase_fcr": "FCR",
            }
        )
        for col in ["Duração estimada (dias)"]:
            if col in growth_view.columns:
                growth_view[col] = pd.to_numeric(growth_view[col], errors="coerce").round(0).astype("Int64")
        for col in ["Ganho médio da fase (g/dia)", "FCR"]:
            if col in growth_view.columns:
                growth_view[col] = pd.to_numeric(growth_view[col], errors="coerce").round(2)
        st.dataframe(growth_view, use_container_width=True, hide_index=True)
    st.caption(
        "Esta tabela confirma que a curva de crescimento está alimentando o cálculo do ciclo. "
        "O peso inicial, peso final, temperatura e o ajuste global alteram a duração das fases e o ciclo total."
    )
    st.markdown('</div>', unsafe_allow_html=True)


with tab1:
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    st.subheader("Sugestão operacional do sistema")
    s1, s2, s3, s4 = st.columns(4)
    with s1:
        st.metric("Ciclo biológico do lote", f"{num(base_res.get('estimated_cycle_days'), 0)} dias")
    with s2:
        st.metric("Primeira receita após", f"{num(base_res.get('production_schedule', {}).get('first_harvest_after_months'), 1)} meses")
    with s3:
        st.metric("Tanques informados", f"{int(fd.get('number_of_units', 0))}")
    with s4:
        st.metric("Tanques mínimos sugeridos", f"{int(base_res.get('production_schedule', {}).get('minimum_units_suggested', fd.get('number_of_units', 1)))}")

    csum1, csum2 = st.columns(2)
    with csum1:
        st.write(f"**Status operacional:** {base_res.get('production_schedule', {}).get('status', '-')}")
        st.write(f"**Período de rampa até a 1ª receita:** {num(base_res.get('production_schedule', {}).get('warmup_months_without_revenue'), 1)} mês(es)")
        st.write(f"**Intervalo real entre despescas em regime:** {num(base_res.get('production_schedule', {}).get('actual_interval_months_with_informed_units'), 2)} mês(es)")
        st.write(f"**Estratégia:** {base_res.get('production_schedule', {}).get('production_strategy', '-')}")
        st.write(f"**Critério:** {base_res.get('production_schedule', {}).get('scheduling_basis', '-')}")
    with csum2:
        st.write(f"**Biomassa por tanque:** {num(base_res.get('final_biomass_per_tank_kg'))} kg")
        st.write(f"**Produção no 1º ano (com rampa):** {num(base_res.get('production_schedule', {}).get('production_year1_kg'))} kg")
        st.write(f"**Peixes colhidos por tanque:** {num(base_res.get('fish_harvested_per_tank'), 0)}")
        st.write(f"**Volume por tanque:** {num(base_res.get('geometry', {}).get('unit_volume_m3'))} m³")
        st.write(f"**Área superficial por tanque:** {num(base_res.get('geometry', {}).get('surface_area_m2'))} m²")
    st.info(base_res.get('production_schedule', {}).get('recommendation', ''))
    if base_res.get("geometry_warning"):
        st.warning(base_res["geometry_warning"])
    st.markdown('</div>', unsafe_allow_html=True)

with tab4:
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    st.subheader("Síntese de aeração")
    a1, a2, a3, a4 = st.columns(4)
    with a1:
        st.metric("Demanda de O₂ / tanque", f"{num(base_res.get('oxygen_demand_per_tank_kg_h'))} kg/h")
    with a2:
        st.metric("Demanda total de O₂", f"{num(base_res.get('oxygen_demand_kg_h'))} kg/h")
    with a3:
        st.metric("Oferta instalada de O₂", f"{num(base_res.get('installed_oxygen_supply_kg_h'))} kg/h")
    with a4:
        st.metric("Equipamentos instalados", f"{int(base_res.get('required_aerators', 0))}")

    if altitude_correction_ui:
        st.caption(
            "Correção por altitude: "
            f"altitude = {num(altitude_correction_ui.get('altitude_m'), 0)} m; "
            f"fator = {num(altitude_correction_ui.get('altitude_factor'), 2)}x; "
            f"oferta equivalente ao nível do mar = {num(altitude_correction_ui.get('installed_supply_original'), 2)} kg/h; "
            f"oferta efetiva corrigida = {num(altitude_correction_ui.get('installed_supply_corrected'), 2)} kg/h; "
            f"relação oferta/demanda = {num(altitude_correction_ui.get('oxygen_offer_demand_ratio'), 2)}x."
        )
        if abs(float(altitude_correction_ui.get("installed_supply_original", 0) or 0) - float(altitude_correction_ui.get("installed_supply_corrected", 0) or 0)) > 0.01:
            st.info(
                "A oferta de oxigênio exibida já está corrigida pela altitude informada. "
                "Quanto maior a altitude, menor a eficiência efetiva de transferência de oxigênio e maior pode ser a necessidade de equipamentos."
            )

    r1, r2 = st.columns(2)
    with r1:
        st.write(f"**Tecnologia adotada:** {base_res.get('selected_aeration_technology', '-')}")
        st.write(f"**Modelo adotado:** {base_res.get('selected_aeration_model', '-')}")
        st.write(f"**Potência instalada total:** {num(base_res.get('peak_installed_power_kw'))} kW")
        st.write(f"**Potência média utilizada no ciclo:** {num(base_res.get('average_active_power_kw'))} kW")
        st.write(f"**Estratégia de operação:** {base_res.get('aeration_operation_mode', '-')}")
    with r2:
        st.write(f"**Estratégia de controle sugerida:** {base_res.get('aeration_control_strategy', '-')}")
        st.write(f"**Equipamento de controle recomendado:** {base_res.get('aeration_control_equipment', '-')}")
        st.write(f"**Custo da aeração (potência fixa):** {brl(base_res.get('aeration_energy_cost_cycle_fixed_power'))}")
        st.write(f"**Custo da aeração (estratégia adotada):** {brl(base_res.get('aeration_energy_cost_cycle'))}")
        st.write(f"**Economia estimada no ciclo:** {brl(base_res.get('aeration_savings_cycle_rs'))} ({num(base_res.get('aeration_savings_cycle_pct'),1)}%)")
    st.caption(base_res.get('aeration_control_note', ''))

    if base_res.get("selected_aeration_warning"):
        st.warning(base_res["selected_aeration_warning"])

    det = base_res.get("aeration_details", {})
    if det:
        st.markdown("**Sugestão técnica de aeração**")
        if isinstance(det, dict) and det.get("technology") == "Soprador":
            st.write(f"- **Família sugerida:** {det.get('blower_family', '-')}")
            st.write(f"- **Modelo sugerido:** {det.get('model', '-')}")
            st.write(f"- **Quantidade total:** {int(det.get('qty_system', 0))}")
            st.write(f"- **Capacidade efetiva por soprador:** {num(det.get('effective_oxygen_capacity_each_kg_h'))} kg O₂/h")
            st.write(f"- **Oferta total instalada:** {num(det.get('installed_supply_total_kg_h'))} kg O₂/h")
            st.write(f"- **Potência unitária:** {num(det.get('power_kw_each'))} kW")
            st.write(f"- **Vazão unitária:** {num(det.get('airflow_m3_h_each'),0)} m³/h")
            st.write(f"- **Pressão nominal:** {num(det.get('pressure_mbar'),0)} mbar")
            diffuser = det.get("diffuser_layout", {}) if isinstance(det, dict) else {}
            if diffuser:
                st.write(f"- **Difusor/mangueira sugerido:** {diffuser.get('diffuser_name', '-')}")
                st.write(f"- **OTR preliminar usado:** {num(diffuser.get('kg_o2_per_m_h'), 3)} kg O₂/h/m ({diffuser.get('otr_reference_class', '-')})")
                st.write(f"- **Metragem calculada por tanque:** {num(diffuser.get('meters_per_tank_required'), 1)} m")
                st.write(f"- **Metragem operacional recomendada por tanque:** {num(diffuser.get('meters_per_tank_recommended'), 1)} m")
                st.write(f"- **Metragem total recomendada:** {num(diffuser.get('meters_total_recommended'), 1)} m")
                st.write(f"- **Anéis/linhas estimados por tanque:** {int(diffuser.get('estimated_rings_per_tank', 0) or 0)}")
                if diffuser.get("is_metrage_excessive"):
                    st.warning(diffuser.get("layout_note", "Metragem acima da faixa operacional/econômica recomendada."))
                st.caption(diffuser.get("diffuser_note", ""))
                st.caption(diffuser.get("density_note", ""))
        elif isinstance(det, dict) and "rows" in det and det["rows"]:
            detalhe_df = pd.DataFrame(det["rows"])
            detalhe_df = detalhe_df.rename(columns={
                "tecnologia": "Tecnologia",
                "modelo": "Modelo",
                "quantidade": "Quantidade total",
                "quantidade_por_tanque": "Quantidade por tanque",
                "arranjo": "Arranjo",
                "arranjo_descricao": "Descrição do arranjo",
                "oferta_por_tanque_kg_h": "Oferta por tanque (kg O₂/h)",
            })
            st.dataframe(detalhe_df, use_container_width=True, hide_index=True)
            if str(det.get("technology", "")).startswith("Híbrido") or str(base_res.get("selected_aeration_technology", "")).startswith("Híbrido"):
                st.info(
                    "Configuração híbrida: o sistema combina tecnologias para separar oxigenação, circulação/vórtice, "
                    "segurança operacional e custo energético. Validar arranjo final em campo."
                )
        elif isinstance(det, dict):
            oferta_tanque = det.get("installed_supply_per_tank_kg_h", 0.0)
            demanda_tanque = base_res.get("oxygen_demand_per_tank_kg_h", 0.0)
            margem = ((oferta_tanque / demanda_tanque) - 1.0) * 100.0 if demanda_tanque > 0 else 0.0
            st.write(f"- **Modelo sugerido:** {det.get('model', '-')}")
            st.write(f"- **Quantidade por tanque:** {int(det.get('qty_per_tank', 0))}")
            st.write(f"- **Quantidade total:** {int(base_res.get('required_aerators', 0))}")
            st.write(f"- **Oferta por tanque:** {num(oferta_tanque)} kg O₂/h")
            st.write(f"- **Demanda por tanque:** {num(demanda_tanque)} kg O₂/h")
            st.write(f"- **Margem de segurança estimada:** {num(margem,1)}%")
            st.write(f"- **Máximo geométrico por tanque:** {int(det.get('max_units_per_tank', 0))}")
            if det.get("arrangement"):
                st.write(f"- **Arranjo geométrico:** {det.get('arrangement')} — {det.get('arrangement_label', '-')}")
                st.write(f"- **Raio de instalação:** {num(det.get('install_radius_m'))} m")
                st.write(f"- **Folga até a borda:** {num(det.get('edge_clearance_m'))} m")
            st.write(f"- **Situação:** {'Viável' if det.get('feasible', False) else 'Inadequada para este arranjo'}")
            if det.get("technology") == "Pás" and fd.get("system_type") == "Circular revestido":
                st.caption("Em tanque circular grande, pás podem ser usadas como apoio à circulação/vórtice; validar posicionamento em campo para concentrar sólidos sem prejudicar o dreno central.")

    aer_phase_rows_ui = _build_aeration_phase_table_rows(base_res, fd)
    if aer_phase_rows_ui:
        st.markdown("**Operação da aeração por fase**")
        aer_phase_view = pd.DataFrame(aer_phase_rows_ui)
        st.dataframe(aer_phase_view, use_container_width=True, hide_index=True)

    st.markdown('</div>', unsafe_allow_html=True)

with tab3:
    phase_summary = pd.DataFrame(base_res.get("feeding_plan", []))
    if not phase_summary.empty:
        st.markdown('<div class="section-box">', unsafe_allow_html=True)
        st.subheader("Duração estimada das fases")
        st.caption("Esses tempos são calculados automaticamente com base nos dados informados no projeto. Não são valores fixos.")

        metric_cols = st.columns(len(phase_summary) + 1)
        for idx, (_, row) in enumerate(phase_summary.iterrows()):
            with metric_cols[idx]:
                st.metric(str(row.get("phase_name", f"Fase {idx+1}")), f"{int(round(row.get('days_in_phase', 0)))} dias")
        with metric_cols[-1]:
            st.metric("Ciclo estimado", f"{int(round(base_res.get('estimated_cycle_days', 0)))} dias")

        view_cols = [
            "phase_name",
            "weight_range",
            "days_in_phase",
            "cumulative_days_end",
            "cycle_share_pct",
            "phase_daily_growth_g",
            "phase_fcr",
        ]
        phase_summary_view = phase_summary[[c for c in view_cols if c in phase_summary.columns]].copy()
        rename_map = {
            "phase_name": "Fase",
            "weight_range": "Faixa de peso",
            "days_in_phase": "Duração estimada (dias)",
            "cumulative_days_end": "Dias acumulados até o fim",
            "cycle_share_pct": "% do ciclo",
            "phase_daily_growth_g": "Ganho médio da fase (g/peixe/dia)",
            "phase_fcr": "FCR da fase",
        }
        phase_summary_view = phase_summary_view.rename(columns=rename_map)
        for col in ["Duração estimada (dias)", "Dias acumulados até o fim"]:
            if col in phase_summary_view.columns:
                phase_summary_view[col] = phase_summary_view[col].round(0).astype(int)
        for col in ["% do ciclo", "Ganho médio da fase (g/peixe/dia)", "FCR da fase"]:
            if col in phase_summary_view.columns:
                phase_summary_view[col] = phase_summary_view[col].round(2)
        st.dataframe(phase_summary_view, use_container_width=True)
        st.info("Use essa tabela como balizador de biometria. Se o peixe real estiver abaixo do peso esperado ao fim de uma fase, revise manejo alimentar, densidade, temperatura, oxigênio e qualidade de água.")
        st.markdown('</div>', unsafe_allow_html=True)

with tab5:
    st.subheader("Indicadores principais")

    save_col1, save_col2 = st.columns([1.2, 2.8])
    with save_col1:
        if st.button("Salvar projeto completo", type="primary", key="btn_salvar_projeto_completo"):
            payload = build_project_payload(
                st.session_state.dash_form_data,
                results,
                st.session_state.current_project_id,
            )
            project_name_for_save = str(st.session_state.dash_form_data.get("project_name", "Projeto sem nome")).strip() or "Projeto sem nome"
            payload["project_name"] = project_name_for_save
            payload["name"] = project_name_for_save
            payload.setdefault("project_meta", {})
            payload["project_meta"]["project_name"] = project_name_for_save
            payload.setdefault("inputs", {})
            payload["inputs"]["project_name"] = project_name_for_save
            try:
                project_id = _save_project_active(payload, st.session_state.current_project_id)
                st.session_state.current_project_id = project_id
                st.session_state.selected_project_id = project_id
                st.session_state.latest_results = results
                st.success(f"Projeto salvo com resultados: {project_id}")
            except Exception as exc:
                st.error(f"Erro ao salvar projeto: {exc}")

    with save_col2:
        if st.session_state.current_project_id:
            st.caption(f"Projeto ativo: {st.session_state.current_project_id}")
        else:
            st.caption("Projeto ainda não salvo. Use o botão ao lado para registrar este cenário.")

    c1, c2, c3, c4 = st.columns(4)

    cards = [
        ("Produção por ciclo biológico", f"{num(base_res.get('production_per_cycle_kg'))} kg"),
        ("Ração/ciclo", f"{num(base_res.get('feed_consumption_cycle_kg'))} kg"),
        ("Ração/ano", f"{num(base_res.get('feed_consumption_year_kg'))} kg"),
        ("Dias do ciclo biológico", num(base_res.get("estimated_cycle_days"), 0)),
    ]

    for col, (title, value) in zip([c1, c2, c3, c4], cards):
        with col:
            st.markdown(
                f"""
<div class="kpi-card">
    <div class="kpi-title">{title}</div>
    <div class="kpi-value">{value}</div>
</div>
""",
                unsafe_allow_html=True,
            )

    c5, c6, c7, c8 = st.columns(4)

    cards2 = [
        ("Margem/ciclo", brl(base_res.get("gross_margin_cycle"))),
        (
            "FCR do plano",
            num(base_res.get("implied_fcr")) if base_res.get("implied_fcr") is not None else "-",
        ),
        ("Custo ração/ciclo", brl(base_res.get("feed_cost_cycle"))),
        ("OPEX/ciclo", brl(base_res.get("opex_cycle"))),
    ]

    for col, (title, value) in zip([c5, c6, c7, c8], cards2):
        with col:
            st.markdown(
                f"""
<div class="kpi-card">
    <div class="kpi-title">{title}</div>
    <div class="kpi-value">{value}</div>
</div>
""",
                unsafe_allow_html=True,
            )

    c9, c10, c11, c12 = st.columns(4)

    cards3 = [
        ("Receita/ciclo", brl(base_res.get("revenue_cycle"))),
        ("Custo alevino/ciclo", brl(base_res.get("fingerling_cost_cycle"))),
        ("Custo aeração/ciclo", brl(base_res.get("aeration_energy_cost_cycle"))),
        ("Aeradores instalados", num(base_res.get("required_aerators"), 0)),
    ]

    for col, (title, value) in zip([c9, c10, c11, c12], cards3):
        with col:
            st.markdown(
                f"""
<div class="kpi-card">
    <div class="kpi-title">{title}</div>
    <div class="kpi-value">{value}</div>
</div>
""",
                unsafe_allow_html=True,
            )

    aeration_info_cols = st.columns(4)
    aeration_info = [
        ("Fator base da aeração", num(temp_factor_ui, 2)),
        ("Fator altitude", num(alt_factor_ui, 2)),
        ("Eficiência de campo", f"{num(fd_calc.get('field_efficiency_pct', 85.0), 0)}%"),
        ("OD alvo", f"{num(fd_calc.get('target_do_mg_l', 5.0), 1)} mg/L"),
    ]
    for col, (title, value) in zip(aeration_info_cols, aeration_info):
        with col:
            st.markdown(
                f"""
<div class="kpi-card">
    <div class="kpi-title">{title}</div>
    <div class="kpi-value">{value}</div>
</div>
""",
                unsafe_allow_html=True,
            )

    schedule = base_res.get("production_schedule", {})
    if schedule:
        st.markdown('<div class="section-box">', unsafe_allow_html=True)
        st.subheader("Estratégia operacional de produção")

        s1, s2, s3, s4 = st.columns(4)

        if schedule.get("production_strategy") == "Escalonada":
            schedule_cards = [
                ("Produção por despesca", f"{num(schedule.get('production_per_harvest_kg'))} kg"),
                ("Tanques por lote", num(schedule.get("units_per_batch_exact"), 2)),
                ("Lotes em paralelo", num(schedule.get("batches_in_parallel"), 0)),
                ("1ª despesca após", f"{num(schedule.get('first_harvest_after_months'), 1)} meses"),
            ]
        else:
            schedule_cards = [
                ("Modo", schedule.get("production_strategy", "-")),
                ("Tanques por ciclo", num(schedule.get("units_per_batch_exact"), 0)),
                ("Colheitas/ano", num(schedule.get("harvests_per_year"), 2)),
                ("Produção por colheita", f"{num(schedule.get('production_per_harvest_kg'))} kg"),
            ]

        for col, (title, value) in zip([s1, s2, s3, s4], schedule_cards):
            with col:
                st.markdown(
                    f"""
<div class="kpi-card">
    <div class="kpi-title">{title}</div>
    <div class="kpi-value">{value}</div>
</div>
""",
                    unsafe_allow_html=True,
                )

        if schedule.get("production_strategy") == "Escalonada":
            mode_text = schedule.get("production_batch_mode", "Automático")
            st.info(
                f"Modo {mode_text}. Com {int(fd['number_of_units'])} tanques e intervalo de {num(schedule.get('harvest_interval_months'), 2)} mês(es), "
                f"o sistema trabalha com {num(schedule.get('batches_in_parallel'), 0)} lotes em paralelo. "
                f"Distribuição sugerida: {schedule.get('batch_distribution_note', '-')}"
            )
            if schedule.get("production_batch_mode") == "Personalizado" and schedule.get("automatic_parallel_batches"):
                st.caption(
                    f"Lotes automáticos estimados pelo tempo de ciclo: {num(schedule.get('automatic_parallel_batches'), 0)}. "
                    "Se o valor personalizado for menor, você estará assumindo um escalonamento mais agressivo que o tempo biológico calculado."
                )
            st.caption(
                "A primeira despesca acontece após o tempo completo de cultivo do primeiro lote; depois disso, a retirada passa a seguir o intervalo definido."
            )
        else:
            st.info("No modo de ciclos simultâneos, todos os tanques entram e saem do ciclo ao mesmo tempo.")

        st.markdown("</div>", unsafe_allow_html=True)

    df_growth = pd.DataFrame(base_res.get("growth_curve", []))
    df_feed = pd.DataFrame(base_res.get("feeding_plan", []))
    df_cost = pd.DataFrame(base_res.get("cost_curve", []))

    if not df_feed.empty:
        st.markdown('<div class="section-box">', unsafe_allow_html=True)
        st.subheader("Tabela do plano alimentar por fase")
        st.caption("Para alterar FCR, use a aba Alimentação avançada.")
        columns = [
            "phase_name",
            "weight_range",
            "days_in_phase",
            "cumulative_days_end",
            "cycle_share_pct",
            "protein_percent",
            "pellet_mm",
            "feed_price_per_kg",
            "phase_fcr",
            "biomass_gain_phase_kg",
            "feed_phase_kg",
            "feed_phase_cost",
        ]
        existing = [c for c in columns if c in df_feed.columns]
        df_feed_view = df_feed[existing].rename(columns={
            "phase_name": "Fase",
            "weight_range": "Faixa de peso",
            "days_in_phase": "Dias da fase",
            "cumulative_days_end": "Dias acumulados",
            "cycle_share_pct": "% do ciclo",
            "protein_percent": "PB (%)",
            "pellet_recommendation": "Granulometria recomendada",
            "pellet_mm": "Pellet base (mm)",
            "feed_price_per_kg": "Preço/kg",
            "phase_fcr": "FCR",
            "biomass_gain_phase_kg": "Ganho de biomassa (kg)",
            "feed_phase_kg": "Ração da fase (kg)",
            "feed_phase_cost": "Custo da fase (R$)",
        })
        for col in ["Dias da fase", "Dias acumulados"]:
            if col in df_feed_view.columns:
                df_feed_view[col] = df_feed_view[col].round(0).astype(int)
        for col in ["% do ciclo", "PB (%)", "Preço/kg", "FCR", "Ganho de biomassa (kg)", "Ração da fase (kg)", "Custo da fase (R$)"]:
            if col in df_feed_view.columns:
                df_feed_view[col] = df_feed_view[col].round(2)
        st.dataframe(df_feed_view, use_container_width=True)
        st.caption("Na fase 1, a granulometria deve evoluir conforme a abertura de boca do peixe. Para 1–20 g, a recomendação prática é trabalhar com faixa progressiva, e não com um único pellet fixo.")
        st.markdown("</div>", unsafe_allow_html=True)

    if not df_growth.empty and "day" in df_growth.columns and "weight_g" in df_growth.columns:
        st.markdown('<div class="section-box">', unsafe_allow_html=True)
        st.subheader("Crescimento ao longo do ciclo")
        fig1 = plt.figure(figsize=(8, 4))
        plt.plot(df_growth["day"], df_growth["weight_g"])
        plt.xlabel("Dia")
        plt.ylabel("Peso (g)")
        plt.tight_layout()
        st.pyplot(fig1)
        st.markdown("</div>", unsafe_allow_html=True)

    if not df_growth.empty and "day" in df_growth.columns and "biomass_kg" in df_growth.columns:
        st.markdown('<div class="section-box">', unsafe_allow_html=True)
        st.subheader("Biomassa ao longo do ciclo")
        fig2 = plt.figure(figsize=(8, 4))
        plt.plot(df_growth["day"], df_growth["biomass_kg"])
        plt.xlabel("Dia")
        plt.ylabel("Biomassa (kg)")
        plt.tight_layout()
        st.pyplot(fig2)
        st.markdown("</div>", unsafe_allow_html=True)

    c13, c14 = st.columns(2)

    with c13:
        if not df_feed.empty and "phase_name" in df_feed.columns and "feed_phase_kg" in df_feed.columns:
            st.markdown('<div class="section-box">', unsafe_allow_html=True)
            st.subheader("Ração por fase")
            fig3 = plt.figure(figsize=(8, 4))
            plt.bar(df_feed["phase_name"], df_feed["feed_phase_kg"])
            plt.xlabel("Fase")
            plt.ylabel("Ração (kg)")
            plt.tight_layout()
            st.pyplot(fig3)
            st.markdown("</div>", unsafe_allow_html=True)

    with c14:
        if not df_feed.empty and "phase_name" in df_feed.columns and "feed_phase_cost" in df_feed.columns:
            st.markdown('<div class="section-box">', unsafe_allow_html=True)
            st.subheader("Custo por fase")
            fig4 = plt.figure(figsize=(8, 4))
            plt.bar(df_feed["phase_name"], df_feed["feed_phase_cost"])
            plt.xlabel("Fase")
            plt.ylabel("Custo (R$)")
            plt.tight_layout()
            st.pyplot(fig4)
            st.markdown("</div>", unsafe_allow_html=True)

    if not df_cost.empty and "day" in df_cost.columns and "cumulative_feed_cost" in df_cost.columns:
        st.markdown('<div class="section-box">', unsafe_allow_html=True)
        st.subheader("Custo acumulado da ração")
        fig5 = plt.figure(figsize=(8, 4))
        plt.plot(df_cost["day"], df_cost["cumulative_feed_cost"])
        plt.xlabel("Dia")
        plt.ylabel("Custo acumulado da ração (R$)")
        plt.tight_layout()
        st.pyplot(fig5)
        st.markdown("</div>", unsafe_allow_html=True)



def _safe_text(value, default: str = "-") -> str:
    """Converte valores para texto seguro, evitando None/campos vazios em relatórios."""
    if value is None:
        return default
    text = str(value).strip()
    return text if text else default


def _safe_number(value, decimals: int = 2, suffix: str = "", default: str = "-") -> str:
    """Formata números com segurança usando o padrão brasileiro já usado no app."""
    try:
        if value is None:
            return default
        return f"{num(float(value), decimals)}{suffix}"
    except Exception:
        return default


def _safe_money(value, default: str = "-") -> str:
    try:
        if value is None:
            return default
        return brl(float(value))
    except Exception:
        return default


def _get_first(mapping: dict, keys: list[str], default=None):
    """Busca o primeiro valor existente em uma lista de possíveis chaves."""
    if not isinstance(mapping, dict):
        return default
    for key in keys:
        if key in mapping and mapping.get(key) not in (None, ""):
            return mapping.get(key)
    return default


def _as_float(value, default: float | None = None):
    try:
        if value is None or value == "":
            return default
        if isinstance(value, str):
            value = value.replace("R$", "").replace("%", "").replace("x", "").strip()
            value = value.replace(".", "").replace(",", ".")
        return float(value)
    except Exception:
        return default


def _markdown_table(headers: list[str], rows: list[list[object]]) -> str:
    """Monta tabela Markdown válida, sem células vazias."""
    clean_headers = [_safe_text(h) for h in headers]
    clean_rows = [[_safe_text(cell) for cell in row] for row in rows]
    lines = [
        "| " + " | ".join(clean_headers) + " |",
        "| " + " | ".join(["---"] * len(clean_headers)) + " |",
    ]
    for row in clean_rows:
        if len(row) < len(clean_headers):
            row = row + ["-"] * (len(clean_headers) - len(row))
        lines.append("| " + " | ".join(row[: len(clean_headers)]) + " |")
    return "\n".join(lines)



def _find_diffuser_layout_any(obj):
    """Localiza layout de mangueira/difusor em resultados simples ou híbridos."""
    if isinstance(obj, dict):
        if isinstance(obj.get("diffuser_layout"), dict):
            return obj.get("diffuser_layout")
        if isinstance(obj.get("layout_hibrido"), dict):
            return obj.get("layout_hibrido")
        for key in ("blower_details", "details", "aeration_details", "suggested_blower_details"):
            found = _find_diffuser_layout_any(obj.get(key))
            if found:
                return found
        rows = obj.get("rows")
        if isinstance(rows, list):
            for item in rows:
                found = _find_diffuser_layout_any(item)
                if found:
                    return found
    elif isinstance(obj, list):
        for item in obj:
            found = _find_diffuser_layout_any(item)
            if found:
                return found
    return None


def _as_int_for_diagram(value, default: int = 0) -> int:
    try:
        if isinstance(value, str):
            if value.lower() in ("sistema", "system"):
                return default
            value = value.replace(",", ".")
        return max(0, int(round(float(value))))
    except Exception:
        return default


def _diagram_has_soprador(rows: list[dict], selected_technology: str) -> bool:
    """Retorna True somente quando a tecnologia efetivamente usa soprador/difusão."""
    tech_text = str(selected_technology or "")
    if "Soprador" in tech_text:
        return True
    for row in rows:
        if "Soprador" in str(row.get("technology", "")):
            return True
    return False


def _equipment_qty_per_tank_for_diagram(item: dict, number_of_units: int) -> tuple[int, int, str]:
    """Interpreta corretamente quantidade total x quantidade por tanque para a planta.

    Nas configurações automáticas/híbridas o cálculo costuma trazer quantidade_por_tanque.
    Na configuração manual, a chave quantidade representa o TOTAL do sistema; portanto
    deve ser dividida pelo número de tanques para desenhar um tanque representativo.
    """
    units = max(1, int(number_of_units or 1))
    per_tank_raw = item.get("quantidade_por_tanque", item.get("qty_per_tank", item.get("quantity_per_tank", None)))
    total_raw = item.get("quantidade", item.get("quantity_total", None))

    if per_tank_raw not in (None, "", "sistema", "system"):
        qpt = _as_int_for_diagram(per_tank_raw, 0)
        total = _as_int_for_diagram(total_raw, qpt * units) if total_raw not in (None, "") else qpt * units
        return qpt, total, "por_tanque"

    total = _as_int_for_diagram(total_raw, 0)
    if total > 0:
        qpt = max(1, int(math.ceil(total / units)))
        return qpt, total, "total_do_sistema"

    return 0, 0, "indefinido"


def _aeration_rows_for_diagram(aeration_details: dict, selected_technology: str, number_of_units: int) -> list[dict]:
    """Normaliza equipamentos para o desenho esquemático do tanque."""
    rows: list[dict] = []
    if isinstance(aeration_details, dict) and isinstance(aeration_details.get("rows"), list):
        for item in aeration_details.get("rows", []):
            if not isinstance(item, dict):
                continue
            qpt, total, qty_origin = _equipment_qty_per_tank_for_diagram(item, number_of_units)
            rows.append({
                "technology": str(item.get("tecnologia", item.get("technology", "-"))),
                "model": str(item.get("modelo", item.get("model", "-"))),
                "qty_per_tank": qpt,
                "quantity_total": total,
                "qty_origin": qty_origin,
                "arrangement": item.get("arranjo", item.get("arrangement", "-")),
                "arrangement_label": item.get("arranjo_descricao", item.get("arrangement_label", "-")),
                "layout_hibrido": item.get("layout_hibrido"),
                "launch_diameter_m": item.get("launch_diameter_m"),
            })
    elif isinstance(aeration_details, dict):
        qpt, total, qty_origin = _equipment_qty_per_tank_for_diagram(aeration_details, number_of_units)
        rows.append({
            "technology": selected_technology,
            "model": str(aeration_details.get("model", aeration_details.get("modelo", "-"))),
            "qty_per_tank": qpt,
            "quantity_total": total,
            "qty_origin": qty_origin,
            "arrangement": aeration_details.get("arrangement", aeration_details.get("arranjo", "-")),
            "arrangement_label": aeration_details.get("arrangement_label", aeration_details.get("arranjo_descricao", "-")),
            "launch_diameter_m": aeration_details.get("launch_diameter_m"),
        })
    return rows


def _paddle_draw_quantity(qty_per_tank: int, model: str = "") -> tuple[int, str]:
    """Quantidade de pás a representar na planta técnica.

    Regra adotada:
    - se o cálculo/inserção indicar 1 pá por tanque, a planta desenha 1 pá;
    - 2 pás opostas entram como alternativa técnica recomendável apenas quando for
      possível substituir 1 modelo mais potente por 2 modelos menores para melhorar
      circulação/vórtice;
    - não força 2 pás quando 1 equipamento menor já atende à demanda.
    """
    qty = max(0, int(qty_per_tank or 0))
    if qty == 1:
        return 1, (
            "Observação: 1 pá por tanque foi mantida na planta porque pode ser suficiente "
            "quando o menor modelo disponível atende à demanda. Quando houver substituição "
            "de 1 equipamento mais potente por 2 menores, recomenda-se avaliar 2 pás "
            "opostas para melhorar vórtice/circulação."
        )
    return qty, ""


def _paddle_angles_and_radius_factor(qty: int) -> tuple[list[float], float]:
    """Padrões da planilha Beraqua/AquaPá: anel externo, fluxo tangencial."""
    qty = max(1, int(qty or 1))
    patterns = {
        1: ([0], 0.65),
        2: ([0, 180], 0.72),
        3: ([0, 120, 240], 0.72),
        4: ([45, 135, 225, 315], 0.72),
        5: ([18, 90, 162, 234, 306], 0.76),
        6: ([0, 60, 120, 180, 240, 300], 0.76),
        7: ([0, 51.4, 102.9, 154.3, 205.7, 257.1, 308.6], 0.76),
        8: ([22.5, 67.5, 112.5, 157.5, 202.5, 247.5, 292.5, 337.5], 0.76),
        9: ([0, 40, 80, 120, 160, 200, 240, 280, 320], 0.80),
        10: ([18, 54, 90, 126, 162, 198, 234, 270, 306, 342], 0.80),
        11: ([0, 32.7, 65.5, 98.2, 130.9, 163.6, 196.4, 229.1, 261.8, 294.5, 327.3], 0.80),
        12: ([15, 45, 75, 105, 135, 165, 195, 225, 255, 285, 315, 345], 0.80),
    }
    if qty in patterns:
        return patterns[qty]
    angles = [360.0 * i / qty for i in range(qty)]
    return angles, 0.80


def _surface_positions(qty: int, radius: float, arrangement: str = "", technology: str = "") -> list[tuple[float, float, float]]:
    """Posições técnicas normalizadas: x, y, ângulo em radianos."""
    qty = max(0, int(qty or 0))
    if qty <= 0:
        return []
    arr = str(arrangement or "").upper()
    tech = str(technology or "")

    if "Pás" in tech or "Pas" in tech or "Pa" in tech:
        draw_qty, _ = _paddle_draw_quantity(qty)
        angles_deg, factor = _paddle_angles_and_radius_factor(draw_qty)
        rr = radius * factor
        # Orientação tangencial para criar corrente circular no mesmo sentido.
        return [(rr * math.cos(math.radians(a)), rr * math.sin(math.radians(a)), math.radians(a + 90.0)) for a in angles_deg]

    if qty == 1:
        return [(0.0, 0.0, 0.0)]

    if "Chafariz" in tech:
        r = radius * 0.62
        if qty == 2 or arr == "A2":
            return [(-r, 0.0, 0.0), (r, 0.0, math.pi)]
        if qty == 3 or arr == "A3":
            return [(r * math.cos(a), r * math.sin(a), a) for a in (math.pi / 2, math.pi / 2 + 2 * math.pi / 3, math.pi / 2 + 4 * math.pi / 3)]
        n = min(qty, 12)
        return [(r * math.cos(2 * math.pi * i / n + math.pi / 4), r * math.sin(2 * math.pi * i / n + math.pi / 4), 2 * math.pi * i / n + math.pi / 4) for i in range(n)]

    n = min(qty, 12)
    r = radius * 0.60
    return [(r * math.cos(2 * math.pi * i / n), r * math.sin(2 * math.pi * i / n), 2 * math.pi * i / n) for i in range(n)]



def _ring_geometry_for_diagram(radius: float, rings: int, target_meters: float | None = None) -> tuple[list[dict], float, int]:
    """Calcula anéis de mangueira/difusor para planta técnica com melhor distribuição radial.

    Regra técnica adotada:
    - a linha externa começa próxima à borda, respeitando afastamento operacional;
    - as demais linhas avançam de forma distribuída até uma zona interna de segurança;
    - a quantidade de anéis é ajustada para representar a metragem recomendada sem
      concentrar todos os anéis apenas junto à parede do tanque;
    - se a metragem recomendada for maior que a geometria comporta, a planta indica
      ramais paralelos por anel.
    """
    try:
        target = float(target_meters) if target_meters is not None else 0.0
    except Exception:
        target = 0.0

    rings_hint = max(0, int(rings or 0))

    # Afastamento operacional: deixa a mangueira próxima à parede, mas não encostada.
    border_clearance = max(0.50, min(0.90, radius * 0.055))
    outer_radius = max(0.80, radius - border_clearance)

    # Zona livre interna para dreno/manutenção e para evitar difusores sobre o dreno.
    inner_clearance = max(0.90, min(2.10, radius * 0.18))
    if inner_clearance >= outer_radius:
        inner_clearance = max(0.35, outer_radius * 0.30)

    # Espaçamento mínimo para manter a planta legível e operacionalmente executável.
    min_spacing = max(0.35, min(0.85, radius * 0.035))

    def build_rows_from_radii(radii: list[float]) -> tuple[list[dict], float]:
        rows: list[dict] = []
        total = 0.0
        for idx, rr in enumerate(radii, start=1):
            diameter = 2.0 * rr
            length = 2.0 * math.pi * rr
            total += length
            spacing = abs(radii[idx - 2] - rr) if idx > 1 else 0.0
            rows.append({
                "idx": idx,
                "radius_m": rr,
                "diameter_m": diameter,
                "length_m": length,
                "spacing_m": spacing,
                "border_clearance_m": radius - outer_radius,
                "inner_clearance_m": inner_clearance,
            })
        return rows, total

    if target > 0:
        best_rows = None
        best_total = 0.0
        best_score = None
        max_n = 24

        # Primeiro: tentar distribuir do anel externo até uma zona interna de segurança.
        # Isso evita o desenho concentrado na borda.
        for n in range(1, max_n + 1):
            if n == 1:
                radii = [outer_radius]
            else:
                spacing = (outer_radius - inner_clearance) / (n - 1)
                if spacing < min_spacing:
                    continue
                radii = [outer_radius - spacing * i for i in range(n)]
            rows, total = build_rows_from_radii(radii)
            relative_error = abs(total - target) / max(target, 1.0)
            hint_penalty = 0.055 * abs(n - rings_hint) if rings_hint > 0 else 0.0
            deficit_penalty = 0.10 if total < target * 0.92 else 0.0
            excess_penalty = 0.04 if total > target * 1.18 else 0.0
            # Preferir desenho que usa a largura útil do tanque e não fica só na borda.
            radial_span = outer_radius - radii[-1]
            span_ratio = radial_span / max(outer_radius - inner_clearance, 0.01)
            span_penalty = max(0.0, 0.95 - span_ratio) * 0.20
            score = (relative_error + hint_penalty + deficit_penalty + excess_penalty + span_penalty, n)
            if best_score is None or score < best_score:
                best_score = score
                best_rows = rows
                best_total = total

        if best_rows is not None:
            # Se a diferença residual for muito grande, informar ramais paralelos por anel.
            multiplier = 1
            if best_total < target * 0.72:
                multiplier = max(1, int(math.ceil(target / max(best_total, 1.0))))
                best_total = best_total * multiplier
            return best_rows, best_total, multiplier

    # Fallback sem metragem recomendada: usar a dica de anéis, mas distribuindo no raio útil.
    rings_final = max(0, rings_hint)
    if rings_final <= 0:
        return [], 0.0, 1
    rings_final = min(rings_final, 18)
    if rings_final == 1:
        radii = [outer_radius]
    else:
        spacing = (outer_radius - inner_clearance) / (rings_final - 1)
        if spacing < min_spacing:
            rings_final = max(1, int((outer_radius - inner_clearance) / min_spacing) + 1)
            rings_final = min(rings_final, 18)
            spacing = (outer_radius - inner_clearance) / max(rings_final - 1, 1)
        radii = [outer_radius - spacing * i for i in range(rings_final)]
    rows, total_geom = build_rows_from_radii(radii)
    return rows, total_geom, 1

def _row_allowed_for_diagram(item: dict, selected_technology: str, form_data: dict) -> bool:
    """Evita que equipamentos residuais de outras escolhas apareçam na planta."""
    tech = str(item.get("technology", ""))
    selected = str(selected_technology or "")
    aeration_mode = str(form_data.get("aeration_mode", form_data.get("aeration_mode_label", "")))
    manual_mode = selected.lower().startswith("configuração manual") or selected.lower().startswith("configuracao manual") or aeration_mode == "Manual"

    if manual_mode:
        if "Chafariz" in tech:
            return bool(form_data.get("manual_use_fountain", False)) and _as_int_for_diagram(form_data.get("manual_fountain_qty"), 0) > 0
        if "Pás" in tech or "Pas" in tech:
            return bool(form_data.get("manual_use_paddlewheel", False)) and _as_int_for_diagram(form_data.get("manual_paddle_qty"), 0) > 0
        if "Soprador" in tech:
            radial_on = bool(form_data.get("manual_use_radial", False)) and _as_int_for_diagram(form_data.get("manual_radial_qty"), 0) > 0
            lobular_on = bool(form_data.get("manual_use_lobular", False)) and _as_int_for_diagram(form_data.get("manual_lobular_qty"), 0) > 0
            return radial_on or lobular_on
        return False

    if selected == "Chafariz":
        return "Chafariz" in tech
    if selected == "Pás":
        return "Pás" in tech or "Pas" in tech
    if selected == "Soprador":
        return "Soprador" in tech
    if selected.startswith("Híbrido") or selected.startswith("Hibrido"):
        if "Chafariz" in tech and "Chafariz" in selected:
            return True
        if ("Pás" in tech or "Pas" in tech) and ("Pás" in selected or "Pas" in selected):
            return True
        if "Soprador" in tech and "Soprador" in selected:
            return True
        return False
    return True


def _create_aeration_layout_diagram(results_data: dict, form_data: dict, output_dir: Path) -> tuple[Path | None, str]:
    """Gera planta técnica do tanque com aeração conforme a configuração calculada."""
    try:
        import matplotlib.pyplot as plt
        from matplotlib.patches import Circle, Rectangle, FancyArrowPatch, Arc
        from matplotlib.lines import Line2D
    except Exception:
        return None, "Matplotlib não disponível para gerar o esquema de aeração."

    try:
        base = results_data.get("base", {}) if isinstance(results_data, dict) else {}
        geometry = base.get("geometry", {}) if isinstance(base, dict) else {}
        details = base.get("aeration_details", {}) if isinstance(base, dict) else {}
        technology = str(base.get("selected_aeration_technology", "-") or "-")
        unit_volume = _as_float(geometry.get("unit_volume_m3"), _as_float(form_data.get("unit_volume_m3"), 0.0)) or 0.0
        depth = _as_float(geometry.get("depth_m"), _as_float(form_data.get("water_depth_m"), 1.2)) or 1.2
        diameter = _as_float(geometry.get("diameter_m"), 0.0) or 0.0
        if diameter <= 0:
            area = _as_float(geometry.get("surface_area_m2"), 0.0) or 0.0
            diameter = math.sqrt((4.0 * area) / math.pi) if area > 0 else 10.0
        radius = max(diameter / 2.0, 1.0)
        surface_area = math.pi * radius * radius

        number_of_units = _as_int_for_diagram(base.get("number_of_units", form_data.get("number_of_units", 1)), 1)
        rows_raw = _aeration_rows_for_diagram(details, technology, number_of_units)
        rows = [r for r in rows_raw if _row_allowed_for_diagram(r, technology, form_data)]
        has_soprador = _diagram_has_soprador(rows, technology)
        diffuser_layout = None
        if has_soprador:
            diffuser_layout = _find_diffuser_layout_any(details) or _find_diffuser_layout_any(base.get("suggested_blower_details", {}))

        color_tank = "#0B1F2A"
        color_diffuser = "#7B2CBF"   # roxo
        color_fountain = "#0077B6"   # azul
        color_paddle = "#2E7D32"     # verde
        color_drain = "#B00020"      # vermelho
        color_dim = "#555555"
        color_flow = "#666666"

        fig = plt.figure(figsize=(14.6, 8.2), dpi=180)
        gs = fig.add_gridspec(1, 2, width_ratios=[2.10, 1.30], wspace=0.08)
        ax = fig.add_subplot(gs[0, 0])
        ax_info = fig.add_subplot(gs[0, 1])
        ax.set_aspect("equal")
        ax.axis("off")
        ax_info.axis("off")

        ax.add_patch(Circle((0, 0), radius, fill=False, linewidth=2.4, edgecolor=color_tank))

        for arc_r in (radius * 0.32, radius * 0.52, radius * 0.72):
            ax.add_patch(Arc((0, 0), 2 * arc_r, 2 * arc_r, theta1=25, theta2=110, linewidth=0.75, color=color_flow, alpha=0.35))
            a = math.radians(110)
            ax.add_patch(FancyArrowPatch((arc_r * math.cos(a - 0.05), arc_r * math.sin(a - 0.05)), (arc_r * math.cos(a), arc_r * math.sin(a)), arrowstyle="->", mutation_scale=7, linewidth=0.75, color=color_flow, alpha=0.45))

        y_dim = -radius * 1.12
        ax.plot([-radius, radius], [y_dim, y_dim], color=color_dim, linewidth=0.9)
        ax.plot([-radius, -radius], [y_dim - radius * 0.025, y_dim + radius * 0.025], color=color_dim, linewidth=0.9)
        ax.plot([radius, radius], [y_dim - radius * 0.025, y_dim + radius * 0.025], color=color_dim, linewidth=0.9)
        ax.text(0, y_dim - radius * 0.065, f"Diâmetro interno estimado: {diameter:.2f} m".replace(".", ","), ha="center", va="top", fontsize=8.4, color=color_dim)

        technical_lines = [
            "DADOS DO TANQUE",
            f"Volume útil: {unit_volume:,.0f} m³".replace(",", "."),
            f"Diâmetro: {diameter:.2f} m".replace(".", ","),
            f"Área superficial: {surface_area:.2f} m²".replace(".", ","),
            f"Lâmina d'água: {depth:.2f} m".replace(".", ","),
            f"Tecnologia: {technology}",
        ]

        ring_rows: list[dict] = []
        total_geom = 0.0
        multiplier = 1
        if isinstance(diffuser_layout, dict):
            rings = _as_int_for_diagram(diffuser_layout.get("estimated_rings_per_tank"), 0)
            meters = _as_float(diffuser_layout.get("meters_per_tank_recommended"), None)
            if rings <= 0 and meters:
                rings = max(1, min(18, int(round(meters / max(1.0, math.pi * radius)))))
            ring_rows, total_geom, multiplier = _ring_geometry_for_diagram(radius, rings, meters)
            for item in ring_rows:
                rr = item["radius_m"]
                ax.add_patch(Circle((0, 0), rr, fill=False, linestyle="--", linewidth=1.25, edgecolor=color_diffuser, alpha=0.92))
                label_angle = math.radians(18 if item["idx"] % 2 else 202)
                lx = rr * math.cos(label_angle)
                ly = rr * math.sin(label_angle)
                ax.text(lx, ly, f"A{item['idx']}\n{item['length_m']:.1f} m".replace(".", ","), fontsize=6.2, ha="center", va="center", color=color_diffuser, bbox=dict(boxstyle="round,pad=0.12", facecolor="white", alpha=0.78, edgecolor="none"))
            if ring_rows:
                spacing = ring_rows[1]["spacing_m"] if len(ring_rows) > 1 else 0.0
                border_distance = ring_rows[0].get("border_clearance_m", radius - ring_rows[0]["radius_m"])
                technical_lines.extend([
                    "",
                    "MANGUEIRA POROSA / DIFUSOR",
                    f"Anéis: {len(ring_rows)} — disposição de fora para dentro",
                    f"Distância do anel externo à borda: {border_distance:.2f} m".replace(".", ","),
                    f"Espaçamento radial entre anéis: {spacing:.2f} m".replace(".", ","),
                    f"Anel interno até o centro/dreno: {ring_rows[-1].get('inner_clearance_m', ring_rows[-1]['radius_m']):.2f} m".replace(".", ","),
                    f"Metragem representada: {total_geom:.1f} m/tanque".replace(".", ","),
                ])
                if meters is not None:
                    technical_lines.append(f"Metragem recomendada pelo cálculo: {meters:.1f} m/tanque".replace(".", ","))
                if multiplier > 1:
                    technical_lines.append(f"Observação: usar ~{multiplier} ramais paralelos por anel para atingir a metragem total.")

        for item in rows:
            tech = str(item.get("technology", ""))
            qpt = _as_int_for_diagram(item.get("qty_per_tank"), 0)
            arrangement = str(item.get("arrangement", ""))
            model = str(item.get("model", "-"))
            if qpt <= 0:
                continue
            if "Chafariz" in tech:
                positions = _surface_positions(qpt, radius, arrangement, tech)
                for idx, (x, y, ang) in enumerate(positions, start=1):
                    launch_d = _as_float(item.get("launch_diameter_m"), 0.0) or radius * 0.28
                    spot_r = max(min(launch_d / 2.0, radius * 0.20), radius * 0.07)
                    ax.add_patch(Circle((x, y), spot_r, fill=False, linestyle=":", linewidth=1.15, edgecolor=color_fountain, alpha=0.86))
                    ax.add_patch(Circle((x, y), radius * 0.043, fill=True, facecolor=color_fountain, edgecolor="white", linewidth=0.9, alpha=0.96))
                    ax.text(x, y + spot_r + radius * 0.035, f"C{idx}", ha="center", va="bottom", fontsize=8, color=color_fountain, fontweight="bold")
                technical_lines.extend(["", f"CHAFARIZ — modelo {model}", f"Quantidade por tanque: {qpt}", f"Arranjo: {item.get('arrangement_label', arrangement)}"])
                if item.get("qty_origin") == "total_do_sistema":
                    technical_lines.append(f"Observação: quantidade original informada: {item.get('quantity_total', 0)} equipamento(s) no sistema.")
            elif "Pás" in tech or "Pas" in tech or "Pa" in tech:
                draw_qpt, paddle_note = _paddle_draw_quantity(qpt, model)
                positions = _surface_positions(draw_qpt, radius, arrangement, tech)
                for idx, (x, y, ang) in enumerate(positions, start=1):
                    w = radius * 0.18
                    h = radius * 0.065
                    rect = Rectangle((x - w / 2, y - h / 2), w, h, angle=math.degrees(ang), fill=True, facecolor=color_paddle, edgecolor="white", linewidth=0.85, alpha=0.90)
                    ax.add_patch(rect)
                    ax.add_patch(FancyArrowPatch((x, y), (x + radius * 0.16 * math.cos(ang), y + radius * 0.16 * math.sin(ang)), arrowstyle="->", mutation_scale=10, linewidth=1.05, color=color_paddle, alpha=0.95))
                    ax.text(x, y + radius * 0.115, f"P{idx}", ha="center", va="bottom", fontsize=8, color=color_paddle, fontweight="bold")
                _, factor = _paddle_angles_and_radius_factor(draw_qpt)
                technical_lines.extend([
                    "",
                    f"AERADOR DE PÁS — modelo {model}",
                    f"Quantidade por tanque representada: {draw_qpt}",
                    f"Raio de instalação: {factor*100:.0f}% do raio útil".replace(".", ","),
                    "Orientação: fluxo tangencial, todos no mesmo sentido",
                ])
                if paddle_note:
                    technical_lines.append(paddle_note)
                if item.get("qty_origin") == "total_do_sistema":
                    technical_lines.append(f"Observação: quantidade original informada: {item.get('quantity_total', 0)} equipamento(s) no sistema.")
            elif "Soprador" in tech:
                technical_lines.extend(["", f"SOPRADOR — modelo {model}", "Distribuição por mangueira/difusor no fundo do tanque"])

        # Dreno central em vermelho, menor que o marcador do chafariz e desenhado por cima.
        ax.add_patch(Circle((0, 0), radius * 0.024, fill=True, facecolor=color_drain, edgecolor="white", linewidth=1.1, alpha=0.99, zorder=12))
        ax.text(0, -radius * 0.070, "Dreno central", ha="center", va="top", fontsize=8.4, color=color_drain, fontweight="bold", zorder=13)

        ax.set_xlim(-radius * 1.18, radius * 1.18)
        ax.set_ylim(-radius * 1.22, radius * 1.18)
        ax.set_title("Planta técnica preliminar — arranjo de oxigenação/aeração por tanque", fontsize=12.5, fontweight="bold", pad=12)

        legend_handles = [
            Line2D([0], [0], color=color_tank, lw=2.2, label="Borda do tanque"),
            Line2D([0], [0], marker="o", color="w", markerfacecolor=color_drain, markersize=6, label="Dreno central"),
        ]
        if ring_rows:
            legend_handles.append(Line2D([0], [0], color=color_diffuser, lw=1.3, linestyle="--", label="Anéis de mangueira/difusor"))
        if any("Chafariz" in str(r.get("technology", "")) for r in rows):
            legend_handles.append(Line2D([0], [0], marker="o", color="w", markerfacecolor=color_fountain, markersize=8, label="Chafariz"))
        if any(("Pás" in str(r.get("technology", "")) or "Pas" in str(r.get("technology", ""))) for r in rows):
            legend_handles.append(Line2D([0], [0], marker="s", color="w", markerfacecolor=color_paddle, markersize=8, label="Aerador de pás"))

        ax_info.legend(handles=legend_handles, loc="upper left", bbox_to_anchor=(0.02, 0.98), fontsize=8.2, frameon=True, framealpha=0.90)
        ax_info.text(
            0.02,
            0.78,
            "\n".join(technical_lines),
            ha="left",
            va="top",
            fontsize=6.8,
            linespacing=0.96,
            color="#102A43",
            bbox=dict(boxstyle="round,pad=0.42", facecolor="white", alpha=0.90, edgecolor="#8FA6B3"),
            transform=ax_info.transAxes,
        )

        if ring_rows:
            table_rows = [[f"A{r['idx']}", f"{r['diameter_m']:.2f}".replace(".", ","), f"{r['length_m']:.1f}".replace(".", ",")] for r in ring_rows[:20]]
            table = ax_info.table(
                cellText=table_rows,
                colLabels=["Anel", "Diâm. (m)", "Comp. (m)"],
                cellLoc="center",
                colLoc="center",
                bbox=[0.02, 0.02, 0.96, 0.30],
            )
            table.auto_set_font_size(False)
            table.set_fontsize(6.5)
            for (row_idx, col_idx), cell in table.get_celld().items():
                cell.set_linewidth(0.45)
                if row_idx == 0:
                    cell.set_text_props(weight="bold")

        safe_project = _slugify_filename(form_data.get("project_name") or "projeto")
        output_dir.mkdir(parents=True, exist_ok=True)
        image_path = output_dir / f"{safe_project}_planta_tecnica_aeracao.png"
        fig.savefig(image_path, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        caption = "Planta técnica preliminar gerada a partir do volume, diâmetro estimado, lâmina d'água e tecnologia selecionada na aba Oxigênio e aeração. Os anéis de mangueira/difusor são representados de fora para dentro, com afastamento da borda e espaçamento radial indicados. A metragem deve ser validada com pressão real do soprador, perda de carga, ficha técnica da mangueira/difusor, hidráulica do tanque, lançamento d'água e segurança elétrica."
        return image_path, caption
    except Exception as exc:
        return None, f"Não foi possível gerar a planta técnica de aeração: {exc}"

def _build_structured_professional_report(results_data: dict, profile: str, form_data: dict) -> str:
    """
    Relatório técnico estruturado e estável.
    Evita placeholders vazios e usa tabelas Markdown bem formadas,
    que depois são convertidas para tabelas reais no DOCX.
    """
    base = results_data.get("base", {}) if isinstance(results_data, dict) else {}
    schedule = base.get("production_schedule", {}) if isinstance(base, dict) else {}
    geometry = base.get("geometry", {}) if isinstance(base, dict) else {}
    feeding_plan = base.get("feeding_plan", []) if isinstance(base, dict) else []
    aer_rows = base.get("aeration_phase_operation", []) if isinstance(base, dict) else []

    def _find_diffuser_layout(obj):
        """Localiza informações de mangueira/difusor em diferentes formatos do cálculo."""
        if isinstance(obj, dict):
            if isinstance(obj.get("diffuser_layout"), dict):
                return obj.get("diffuser_layout")
            if isinstance(obj.get("layout_hibrido"), dict):
                return obj.get("layout_hibrido")
            for key in ("blower_details", "details", "aeration_details"):
                found = _find_diffuser_layout(obj.get(key))
                if found:
                    return found
            rows = obj.get("rows")
            if isinstance(rows, list):
                for item in rows:
                    found = _find_diffuser_layout(item)
                    if found:
                        return found
        elif isinstance(obj, list):
            for item in obj:
                found = _find_diffuser_layout(item)
                if found:
                    return found
        return None

    diffuser_layout = _find_diffuser_layout(base.get("aeration_details", {}))
    if not diffuser_layout:
        diffuser_layout = _find_diffuser_layout(base.get("suggested_blower_details", {}))

    installed_supply = _as_float(base.get("installed_oxygen_supply_kg_h"), 0.0) or 0.0
    oxygen_demand_total = _as_float(base.get("oxygen_demand_kg_h"), 0.0) or 0.0
    offer_demand_ratio = installed_supply / oxygen_demand_total if oxygen_demand_total > 0 else None

    critical_row = None
    critical_score = -1.0
    for row in aer_rows if isinstance(aer_rows, list) else []:
        if not isinstance(row, dict):
            continue
        demand = _as_float(_get_first(row, ["Demanda O₂ (kg/h)", "Demanda O2 (kg/h)", "oxygen_demand_kg_h", "demand_kg_h", "phase_oxygen_demand_kg_h"]), 0.0) or 0.0
        usage = _as_float(_get_first(row, ["% de uso da capacidade", "% de modulação sugerida", "Modulação sugerida", "usage_pct", "capacity_usage_pct"]), None)
        score = usage if usage is not None else demand
        if score > critical_score:
            critical_score = score
            critical_row = row

    if critical_row:
        critical_usage = _as_float(_get_first(critical_row, ["% de uso da capacidade", "% de modulação sugerida", "Modulação sugerida", "usage_pct", "capacity_usage_pct"]), None)
    else:
        critical_usage = None
    if critical_usage is None and offer_demand_ratio and offer_demand_ratio > 0:
        critical_usage = min(100.0, (1.0 / offer_demand_ratio) * 100.0)

    compatibility_warning = _safe_text(base.get("selected_aeration_warning"), "")
    if compatibility_warning:
        compatibility_text = compatibility_warning
    elif installed_supply > 0 and oxygen_demand_total > 0 and installed_supply >= oxygen_demand_total:
        compatibility_text = "Compatível para o cenário simulado, considerando demanda total estimada e oferta instalada de oxigênio."
    elif oxygen_demand_total > 0:
        compatibility_text = "Revisar dimensionamento: a oferta instalada estimada está abaixo da demanda total de oxigênio."
    else:
        compatibility_text = "Compatibilidade dependente da validação operacional do equipamento, geometria dos tanques e manejo de campo."

    project_rows = [
        ["Projeto", form_data.get("project_name")],
        ["Responsável técnico / autor", form_data.get("author_name")],
        ["Perfil do relatório", profile],
        ["Espécie", "Tilápia"],
        ["Região de referência", form_data.get("region_focus")],
        ["Tipo de estrutura", form_data.get("system_type")],
    ]

    production_rows = [
        ["Estratégia de produção", schedule.get("production_strategy", form_data.get("production_strategy"))],
        ["Critério de escalonamento", schedule.get("scheduling_basis", form_data.get("scheduling_basis"))],
        ["Ciclo biológico do lote", f"{_safe_number(base.get('estimated_cycle_days'), 0)} dias"],
        ["Primeira receita após", f"{_safe_number(schedule.get('first_harvest_after_months'), 1)} meses"],
        ["Intervalo entre despescas em regime", f"{_safe_number(schedule.get('actual_interval_months_with_informed_units'), 2)} meses"],
        ["Lotes em paralelo", _safe_number(schedule.get("batches_in_parallel"), 0)],
        ["Tanques por lote", _safe_number(schedule.get("units_per_batch_exact"), 2)],
        ["Produção por despesca em regime", f"{_safe_number(schedule.get('production_per_harvest_kg'), 2)} kg"],
        ["Despescas no 1º ano (com rampa)", _safe_number(schedule.get("harvests_year1"), 0)],
        ["Produção no 1º ano com rampa", f"{_safe_number(schedule.get('production_year1_kg'), 2)} kg"],
        ["Receita no 1º ano com rampa", _safe_money(schedule.get("revenue_year1"))],
        ["Despescas anuais em regime estabilizado", _safe_number(schedule.get("harvests_year2_stable", schedule.get("harvests_per_year")), 2)],
        ["Produção anual em regime estabilizado", f"{_safe_number(schedule.get('production_year2_kg', base.get('production_per_year_kg')), 2)} kg"],
        ["Receita anual em regime estabilizado", _safe_money(schedule.get("revenue_year2", base.get("revenue_year")))],
        ["Produção média mensal em regime", f"{_safe_number(schedule.get('monthly_production_stable_kg'), 2)} kg/mês"],
        ["Receita média mensal em regime", _safe_money(schedule.get("monthly_revenue_stable"))],
        ["Alevinos estocados por lote/despesca", _safe_number(schedule.get("fish_stocked_per_harvest_batch"), 0)],
    ]

    dimension_rows = [
        ["Tanques informados", _safe_number(form_data.get("number_of_units"), 0)],
        ["Volume útil por tanque", f"{_safe_number(geometry.get('unit_volume_m3'), 2)} m³"],
        ["Área superficial por tanque", f"{_safe_number(geometry.get('surface_area_m2'), 2)} m²"],
        ["Densidade final", f"{_safe_number(form_data.get('density_kg_m3'), 2)} kg/m³"],
        ["Sobrevivência estimada", _safe_number(form_data.get("survival_rate"), 2)],
        ["Alevinos estocados no sistema", _safe_number(base.get("fish_stocked"), 0)],
        ["Alevinos estocados por tanque", _safe_number(base.get("fish_stocked_per_tank"), 0)],
        ["Peixes colhidos por tanque", _safe_number(base.get("fish_harvested_per_tank"), 0)],
        ["Biomassa final por tanque", f"{_safe_number(base.get('final_biomass_per_tank_kg'), 2)} kg"],
    ]

    aeration_summary_rows = [
        ["Demanda de oxigênio por tanque", f"{_safe_number(base.get('oxygen_demand_per_tank_kg_h'), 2)} kg O₂/h"],
        ["Demanda total de oxigênio", f"{_safe_number(base.get('oxygen_demand_kg_h'), 2)} kg O₂/h"],
        ["Tecnologia adotada", base.get("selected_aeration_technology")],
        ["Modelo adotado", base.get("selected_aeration_model")],
        ["Equipamentos instalados", _safe_number(base.get("required_aerators"), 0)],
        ["Oferta instalada de oxigênio", f"{_safe_number(base.get('installed_oxygen_supply_kg_h'), 2)} kg O₂/h"],
        ["Potência instalada total", f"{_safe_number(base.get('peak_installed_power_kw'), 2)} kW"],
        ["Potência média utilizada no ciclo", f"{_safe_number(base.get('average_active_power_kw'), 2)} kW"],
        ["Modo de operação", base.get("aeration_operation_mode")],
        ["Estratégia de controle sugerida", base.get("aeration_control_strategy")],
        ["Equipamento de controle recomendado", base.get("aeration_control_equipment")],
        ["Custo de aeração no ciclo (potência fixa)", _safe_money(base.get("aeration_energy_cost_cycle_fixed_power"))],
        ["Custo de aeração no ciclo (estratégia adotada)", _safe_money(base.get("aeration_energy_cost_cycle"))],
        ["Economia estimada no ciclo", f"{_safe_money(base.get('aeration_savings_cycle_rs'))} ({_safe_number(base.get('aeration_savings_cycle_pct'), 1)}%)"],
        ["Uso estimado do modelo na fase crítica", f"{_safe_number(critical_usage, 1)}%"],
        ["Relação oferta/demanda na fase crítica", f"{_safe_number(offer_demand_ratio, 2)}x"],
        ["Compatibilidade geométrica / operacional", compatibility_text],
    ]

    if isinstance(diffuser_layout, dict):
        aeration_summary_rows.extend([
            ["Difusor/mangueira estimado", _safe_text(diffuser_layout.get("diffuser_name"), "-")],
            ["Classe OTR usada para difusão", _safe_text(diffuser_layout.get("otr_reference_class"), "-")],
            ["OTR preliminar usado", f"{_safe_number(diffuser_layout.get('kg_o2_per_m_h'), 3)} kg O₂/h/m"],
            ["Metragem calculada por tanque", f"{_safe_number(diffuser_layout.get('meters_per_tank_required'), 1)} m"],
            ["Metragem recomendada por tanque", f"{_safe_number(diffuser_layout.get('meters_per_tank_recommended'), 1)} m"],
            ["Metragem total recomendada", f"{_safe_number(diffuser_layout.get('meters_total_recommended'), 1)} m"],
            ["Linhas/anéis estimados por tanque", _safe_number(diffuser_layout.get("estimated_rings_per_tank"), 0)],
            ["Observação da difusão", _safe_text(diffuser_layout.get("density_note", diffuser_layout.get("layout_note")), "-")],
        ])

    aeration_phase_dict_rows = _build_aeration_phase_table_rows(base, form_data)
    aeration_phase_rows = [
        [
            row.get("Fase", "-"),
            row.get("Biomassa média (kg)", "-"),
            row.get("Demanda O₂ (kg/h)", "-"),
            row.get("Modulação sugerida", "-"),
            row.get("Equipamentos operando", "-"),
            row.get("Potência média ativa (kW)", "-"),
            row.get("Dias", "-"),
            row.get("Custo da fase", "-"),
        ]
        for row in aeration_phase_dict_rows
    ]

    if not aeration_phase_rows:
        aeration_phase_rows.append(["-", "-", "-", "-", "-", "-", "-", "-"])

    feeding_rows = []
    for idx, row in enumerate(feeding_plan if isinstance(feeding_plan, list) else [], start=1):
        if not isinstance(row, dict):
            continue
        feeding_rows.append([
            row.get("phase_name", f"Fase {idx}"),
            row.get("weight_range", "-"),
            _safe_number(row.get("days_in_phase"), 0),
            _safe_number(row.get("protein_percent"), 1),
            _safe_text(row.get("pellet_recommendation", row.get("pellet_mm", "-"))),
            _safe_money(row.get("feed_price_per_kg")),
            _safe_number(row.get("phase_fcr"), 2),
            _safe_number(row.get("biomass_gain_phase_kg"), 2),
            _safe_number(row.get("feed_phase_kg"), 2),
            _safe_money(row.get("feed_phase_cost")),
        ])

    if not feeding_rows:
        feeding_rows.append(["-", "-", "-", "-", "-", "-", "-", "-", "-", "-"])

    cost_rows = [
        ["Receita por ciclo biológico completo", _safe_money(base.get("revenue_cycle"))],
        ["Custo de alevinos por ciclo", _safe_money(base.get("fingerling_cost_cycle"))],
        ["Custo de ração por ciclo", _safe_money(base.get("feed_cost_cycle"))],
        ["Custo de aeração por ciclo", _safe_money(base.get("aeration_energy_cost_cycle"))],
        ["OPEX por ciclo", _safe_money(base.get("opex_cycle"))],
        ["Margem bruta por ciclo", _safe_money(base.get("gross_margin_cycle"))],
        ["Receita no 1º ano com rampa", _safe_money(schedule.get("revenue_year1"))],
        ["Produção no 1º ano com rampa", f"{_safe_number(schedule.get('production_year1_kg'), 2)} kg"],
        ["Receita anual em regime estabilizado", _safe_money(schedule.get("revenue_year2", base.get("revenue_year")))],
        ["OPEX anual em regime estabilizado", _safe_money(base.get("opex_year"))],
        ["Margem bruta anual em regime estabilizado", _safe_money(base.get("gross_margin_year"))],
        ["Receita média mensal em regime", _safe_money(schedule.get("monthly_revenue_stable"))],
        ["Produção média mensal em regime", f"{_safe_number(schedule.get('monthly_production_stable_kg'), 2)} kg/mês"],
        ["OPEX médio mensal em regime", _safe_money((_as_float(base.get("opex_year"), 0.0) or 0.0) / 12.0)],
        ["Margem bruta média mensal em regime", _safe_money((_as_float(base.get("gross_margin_year"), 0.0) or 0.0) / 12.0)],
        ["FCR do plano", _safe_number(base.get("implied_fcr"), 2)],
    ]

    lines = [
        "# PROJETO TÉCNICO-ECONÔMICO DE PISCICULTURA",
        "",
        "## 1. Identificação do projeto",
        "",
        _markdown_table(["Campo", "Informação"], project_rows),
        "",
        "## 2. Resumo executivo",
        "",
        f"O projeto simula a produção de tilápia em sistema {_safe_text(form_data.get('system_type')).lower()}, com {_safe_text(form_data.get('number_of_units'))} tanque(s) informado(s), peso inicial de {_safe_number(form_data.get('initial_weight_g'), 2)} g e peso final alvo de {_safe_number(form_data.get('target_weight_g'), 2)} g.",
        "",
        "## 3. Estratégia de produção e escalonamento",
        "",
        _markdown_table(["Indicador", "Valor"], production_rows),
        "",
        _safe_text(schedule.get("batch_distribution_note"), "A distribuição dos lotes deve ser validada conforme o cronograma operacional e a disponibilidade de tanques."),
        "",
        "## 4. Dimensionamento do sistema",
        "",
        _markdown_table(["Parâmetro", "Valor"], dimension_rows),
        "",
        "## 5. Oxigênio e aeração",
        "",
        _markdown_table(["Indicador", "Valor"], aeration_summary_rows),
        "",
        "### 5.1 Arranjo esquemático da aeração por tanque",
        "",
        "[[AERATION_LAYOUT_IMAGE]]",
        "",
        "O esquema é ilustrativo e acompanha a tecnologia selecionada na aba Oxigênio e aeração. A validação final deve considerar fixação, lançamento de água, circulação, vórtice, drenagem central, segurança elétrica e acesso para manutenção.",
        "",
        "### 5.2 Operação da aeração por fase",
        "",
        _markdown_table(
            ["Fase", "Biomassa média (kg)", "Demanda O₂ (kg/h)", "Modulação sugerida", "Equipamentos operando", "Potência média ativa (kW)", "Dias", "Custo da fase"],
            aeration_phase_rows,
        ),
        "",
        "## 6. Plano alimentar por fase",
        "",
        _markdown_table(
            ["Fase", "Faixa de peso", "Dias", "PB (%)", "Pellet recomendado", "Preço/kg", "FCR", "Ganho biomassa (kg)", "Ração fase (kg)", "Custo fase"],
            feeding_rows,
        ),
        "",
        "## 7. Custos e indicadores econômicos",
        "",
        _markdown_table(["Indicador", "Valor"], cost_rows),
        "",
        "## 8. Interpretação técnica",
        "",
        "O resultado deve ser interpretado como simulação técnico-econômica. No sistema escalonado, o 1º ano inclui a rampa até a primeira despesca; os indicadores mensais representam o regime estabilizado após a entrada regular das receitas. A validação final depende de biometria real, qualidade de água, desempenho alimentar, manejo sanitário, disponibilidade energética e padronização dos lotes.",
        "",
        "## 9. Recomendações",
        "",
        "- Monitorar oxigênio dissolvido, temperatura, pH, alcalinidade, amônia e transparência/qualidade da água.",
        "- Revisar FCR por fase com base em biometria e consumo real.",
        "- Conferir a compatibilidade geométrica dos aeradores com o tanque antes da compra definitiva.",
        "- Reavaliar preço de ração, energia e venda periodicamente.",
        "- Em produção escalonada, manter controle rigoroso dos lotes para evitar mistura de idades e pesos.",
        "",
        "## 10. Observações finais",
        "",
        _safe_text(form_data.get("notes"), "Projeto gerado pelo Aqua Project Agent Dashboard."),
        "",
    ]
    return "\n".join(lines)


def _ensure_report_content(report_text: str, results_data: dict, profile: str, form_data: dict) -> str:
    """Mantido para compatibilidade; agora sempre retorna relatório estruturado."""
    return _build_structured_professional_report(results_data, profile, form_data)


def _clean_markdown_inline(text: str) -> str:
    """Remove marcações simples de Markdown para o DOCX ficar limpo."""
    clean = str(text)
    for token in ("**", "__", "`"):
        clean = clean.replace(token, "")
    return clean.replace("<br>", " ").replace("<br/>", " ").replace("<br />", " ").strip()


def _is_markdown_table_separator(cells: list[str]) -> bool:
    if not cells:
        return False
    for cell in cells:
        normalized = cell.replace(":", "").replace("-", "").strip()
        if normalized:
            return False
    return True


def _split_markdown_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [_clean_markdown_inline(cell.strip()) for cell in stripped.split("|")]


def _add_docx_table(document, table_lines: list[str]) -> None:
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    parsed_rows = [_split_markdown_table_row(line) for line in table_lines if line.strip().startswith("|")]
    parsed_rows = [row for row in parsed_rows if not _is_markdown_table_separator(row)]
    if not parsed_rows:
        return

    max_cols = max(len(row) for row in parsed_rows)
    parsed_rows = [row + [""] * (max_cols - len(row)) for row in parsed_rows]

    table = document.add_table(rows=1, cols=max_cols)
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for idx, value in enumerate(parsed_rows[0]):
        header_cells[idx].text = value
        for paragraph in header_cells[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8.5)

    for row in parsed_rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    document.add_paragraph("")


def _markdown_to_docx_bytes(markdown_text: str, title: str, author: str, aeration_diagram_path: Path | None = None, aeration_diagram_caption: str = "") -> bytes:
    """Gera DOCX em memória, convertendo tabelas Markdown em tabelas reais do Word."""
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.section import WD_ORIENT

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)

    if title:
        document.add_heading(str(title), level=0)
    if author:
        document.add_paragraph(f"Responsável técnico / autor: {author}")
    document.add_paragraph("")

    lines = (markdown_text or "").splitlines()
    idx = 0
    while idx < len(lines):
        raw_line = lines[idx]
        stripped = raw_line.strip()

        if not stripped:
            idx += 1
            continue

        if stripped == "[[AERATION_LAYOUT_IMAGE]]":
            if aeration_diagram_path and Path(aeration_diagram_path).exists():
                try:
                    document.add_picture(str(aeration_diagram_path), width=Inches(5.8))
                    if aeration_diagram_caption:
                        cap = document.add_paragraph(_clean_markdown_inline(aeration_diagram_caption))
                        for run in cap.runs:
                            run.italic = True
                            run.font.size = Pt(8)
                except Exception as exc:
                    document.add_paragraph(f"Esquema de aeração não inserido no DOCX: {exc}")
            else:
                document.add_paragraph("Esquema de aeração não disponível para esta simulação.")
            idx += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            _add_docx_table(document, table_lines)
            continue

        if stripped.startswith("### "):
            document.add_heading(_clean_markdown_inline(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            document.add_heading(_clean_markdown_inline(stripped[3:]), level=2)
        elif stripped.startswith("# "):
            document.add_heading(_clean_markdown_inline(stripped[2:]), level=1)
        elif stripped.startswith("- "):
            document.add_paragraph(_clean_markdown_inline(stripped[2:]), style="List Bullet")
        elif stripped.startswith("* "):
            document.add_paragraph(_clean_markdown_inline(stripped[2:]), style="List Bullet")
        else:
            document.add_paragraph(_clean_markdown_inline(stripped))
        idx += 1

    bio = BytesIO()
    document.save(bio)
    return bio.getvalue()


with tab6:
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    st.subheader("Relatório Profissional")
    st.caption("Estruturado para produtor, técnico ou banco/financiamento.")

    project_root = Path(__file__).resolve().parent
    save_dir = project_root / output_dir
    save_dir.mkdir(parents=True, exist_ok=True)

    aeration_diagram_path, aeration_diagram_caption = _create_aeration_layout_diagram(results, fd, save_dir)

    professional_report = _build_structured_professional_report(
        results,
        report_profile,
        fd,
    )

    st.markdown("### Pré-visualização do relatório")
    preview_report = professional_report.replace(
        "[[AERATION_LAYOUT_IMAGE]]",
        "_Esquema de aeração inserido no DOCX; prévia exibida abaixo._",
    )
    st.markdown(preview_report)
    if aeration_diagram_path and Path(aeration_diagram_path).exists():
        st.image(str(aeration_diagram_path), caption=aeration_diagram_caption, use_container_width=True)
    elif aeration_diagram_caption:
        st.warning(aeration_diagram_caption)

    st.info(f"Pasta de saída atual: {save_dir}")

    safe_output_name = _slugify_filename(fd.get("project_name") or output_name)
    md_path = save_dir / f"{safe_output_name}_completo.md"
    docx_path = save_dir / f"{safe_output_name}_completo.docx"

    md_bytes = professional_report.encode("utf-8")

    try:
        md_path.write_text(professional_report, encoding="utf-8")
    except Exception as exc:
        st.warning(f"Não foi possível salvar o Markdown na pasta de saída: {exc}")

    docx_bytes = None
    docx_error = None

    try:
        docx_bytes = _markdown_to_docx_bytes(
            professional_report,
            fd.get("project_name", inp.project_name),
            fd.get("author_name", inp.author_name),
            aeration_diagram_path,
            aeration_diagram_caption,
        )
        docx_path.write_bytes(docx_bytes)
    except Exception as exc:
        docx_error = exc

    col1, col2, col3 = st.columns(3)

    with col1:
        st.download_button(
            label="Baixar Relatório (Markdown)",
            data=md_bytes,
            file_name=f"{safe_output_name}_completo.md",
            mime="text/markdown",
            key="download_relatorio_markdown",
        )

    with col2:
        if docx_bytes and len(docx_bytes) > 1000:
            st.download_button(
                label="Baixar Relatório (DOCX)",
                data=docx_bytes,
                file_name=f"{safe_output_name}_completo.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_relatorio_docx",
            )
        else:
            st.error(f"DOCX não foi gerado corretamente: {docx_error}")

    with col3:
        if st.button("Tentar gerar PDF", key="btn_tentar_gerar_pdf"):
            if not docx_path.exists() or docx_path.stat().st_size == 0:
                st.error("O DOCX ainda não foi gerado corretamente; não foi possível tentar converter para PDF.")
            else:
                try:
                    pdf_path = export_docx_to_pdf(docx_path, save_dir)
                    if pdf_path and Path(pdf_path).exists() and Path(pdf_path).stat().st_size > 0:
                        st.session_state["latest_pdf_path"] = str(pdf_path)
                        st.success("PDF gerado. O botão de download aparecerá abaixo.")
                    else:
                        st.warning("A conversão para PDF não retornou um arquivo válido.")
                except Exception as exc:
                    st.warning(
                        "DOCX gerado, mas o PDF não foi gerado neste ambiente. "
                        f"No Streamlit Cloud isso pode acontecer por falta do conversor do Word/LibreOffice. Detalhe: {exc}"
                    )

    latest_pdf_path = st.session_state.get("latest_pdf_path")
    if latest_pdf_path and Path(latest_pdf_path).exists() and Path(latest_pdf_path).stat().st_size > 0:
        pdf_path = Path(latest_pdf_path)
        st.download_button(
            label="Baixar Relatório (PDF)",
            data=pdf_path.read_bytes(),
            file_name=pdf_path.name,
            mime="application/pdf",
            key="download_relatorio_pdf",
        )

    st.caption(
        "Observação: no Streamlit Cloud, o caminho exibido em 'Pasta de saída' é uma pasta do servidor. "
        "Para receber o arquivo no seu computador ou celular, use os botões 'Baixar'."
    )

    st.markdown("</div>", unsafe_allow_html=True)
